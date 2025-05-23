import os
import tkinter as tk
from tkinter import filedialog, ttk, messagebox, simpledialog
import speech_recognition as sr
import pyaudio
import json
import wave
import threading
import tempfile
import time
import math
import sys
from datetime import datetime

import platform

try:
    import sounddevice as sd
    sounddevice_disponible = True
except ImportError:
    sounddevice_disponible = False
    print("Nota: sounddevice no está instalado. La captura de audio del sistema puede no estar disponible.")

try:
    from vosk import Model, KaldiRecognizer
    vosk_disponible = True
except ImportError:
    vosk_disponible = False
    print("Nota: Vosk no está instalado. El modo offline no estará disponible.")

try:
    import shutil
    shutil_disponible = True
except ImportError:
    shutil_disponible = False
    print("Nota: shutil no está instalado.")

try:
    from cryptography.fernet import Fernet
    from cryptography.hazmat.primitives import hashes
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    import base64
    encriptacion_disponible = True
except ImportError:
    encriptacion_disponible = False
    print("Nota: Cryptography no está instalada. La encriptación no estará disponible.")

try:
    from pydub import AudioSegment
    pydub_disponible = True
except ImportError:
    pydub_disponible = False
    print("Nota: pydub no está instalada. La conversión de formatos no estará disponible.")

# Intentar importar la librería docx para soporte de Word
try:
    import docx
    docx_disponible = True
except ImportError:
    docx_disponible = False
    print("Nota: python-docx no está instalada. La exportación a Word no estará disponible.")

class TranscriptorMultilingue:

    def __init__(self, root):
        self.root = root
        self.root.title("Transcriptor Multilingüe con Privacidad")
        self.root.geometry("900x700")  # Tamaño inicial más adecuado

        # Configurar logger
        self.setup_logger()

        # Inicializar variables
        self.audio = pyaudio.PyAudio()
        self.stream = None  # Inicializar stream como None
        self.frames = []  # Inicializar frames para almacenar los datos de audio
        self.modelos_vosk = {}
        self.palabras_clave = set()
        self.estado_var = tk.StringVar()
        self.idioma_var = tk.StringVar(value="es-ES")
        self.modo_reconocimiento = tk.StringVar(value="online")
        self.modo_privado = tk.BooleanVar(value=False)
        self.usar_encriptacion = tk.BooleanVar(value=False)
        self.formato_salida = tk.StringVar(value="txt")
        self.dispositivo_idx = tk.IntVar(value=0)
        self.duracion_segmento = 3
        self.superposicion = 1
        self.reconocedor = sr.Recognizer()
        self.cancelacion_ruido_var = tk.BooleanVar(value=True)
        self.grabacion_activa = False  # Variable para controlar el estado de la grabación
        self.actualizando_tiempo = False  # Inicializar variable para tiempo
        self.tasa_muestreo = 16000  # Establecer tasa de muestreo predeterminada
        self.clave_encriptacion = None
        self.salt = None
        self.contenido_encriptado = None
        self.tiempo_var = tk.StringVar(value="00:00:00")
        self.tamano_var = tk.StringVar(value="0 MB")
        self.palabras_detectadas_var = tk.StringVar(value="0")
        self.grabando = False  # Variable para toggle_grabacion

        # Cargar configuración
        self.cargar_modelos_vosk()
        self.cargar_configuracion_inicial()

        # Crear interfaz
        self.crear_interfaz()

    def setup_logger(self):
        """Configura el logger para registrar eventos"""
        import logging
        from logging.handlers import RotatingFileHandler

        # Crear directorio de logs si no existe
        log_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
        if not os.path.exists(log_dir):
            os.makedirs(log_dir)

        # Configurar logger
        self.logger = logging.getLogger("TranscriptorLog")
        self.logger.setLevel(logging.INFO)

        # Crear handler para archivo con rotación
        log_file = os.path.join(log_dir, f"transcriptor_{datetime.now().strftime('%Y%m%d')}.log")
        handler = RotatingFileHandler(log_file, maxBytes=5*1024*1024, backupCount=3)

        # Formato
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        handler.setFormatter(formatter)

        # Agregar handler al logger
        self.logger.addHandler(handler)

        self.logger.info("=== Inicio de sesión ===")

    def cargar_configuracion_inicial(self):
        """Carga la configuración inicial desde archivo si existe"""
        archivo_config = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")

        if os.path.exists(archivo_config):
            try:
                with open(archivo_config, 'r', encoding='utf-8') as f:
                    config = json.load(f)

                # Aplicar configuración
                if "duracion_segmento" in config:
                    self.duracion_segmento = config["duracion_segmento"]

                if "superposicion" in config:
                    self.superposicion = config["superposicion"]

                if "energy_threshold" in config:
                    self.reconocedor.energy_threshold = config["energy_threshold"]

                if "pause_threshold" in config:
                    self.reconocedor.pause_threshold = config["pause_threshold"]

                if "dynamic_energy_threshold" in config:
                    self.reconocedor.dynamic_energy_threshold = config["dynamic_energy_threshold"]

                if "cancelacion_ruido" in config:
                    self.cancelacion_ruido_var.set(config["cancelacion_ruido"] and pydub_disponible)

                if "formato_salida" in config:
                    self.formato_salida.set(config["formato_salida"])

                if "dispositivo_idx" in config:
                    self.dispositivo_idx.set(config["dispositivo_idx"])

                if "modo_privado" in config:
                    self.modo_privado.set(config["modo_privado"])

                if "usar_encriptacion" in config:
                    self.usar_encriptacion.set(config["usar_encriptacion"])

                if "modo_reconocimiento" in config:
                    self.modo_reconocimiento.set(config["modo_reconocimiento"])

                if "idioma" in config:
                    self.idioma_var.set(config["idioma"])

                self.logger.info("Configuración cargada desde archivo")

            except Exception as e:
                print(f"Error al cargar configuración: {e}")
                self.logger.error(f"Error al cargar configuración: {str(e)}")

        # Cargar palabras clave
        archivo_palabras_clave = os.path.join(os.path.dirname(os.path.abspath(__file__)), "palabras_clave.txt")

        if os.path.exists(archivo_palabras_clave):
            try:
                with open(archivo_palabras_clave, 'r', encoding='utf-8') as f:
                    palabras = [l.strip().lower() for l in f.readlines() if l.strip()]
                    self.palabras_clave = set(palabras)

                self.logger.info(f"Palabras clave cargadas: {len(self.palabras_clave)}")

            except Exception as e:
                print(f"Error al cargar palabras clave: {e}")
                self.logger.error(f"Error al cargar palabras clave: {str(e)}")

    def capturar_audio_sistema(self):
        """Configura la captura del audio del sistema en Windows"""
        try:
            if not sounddevice_disponible:
                messagebox.showerror("Error", 
                                "Para capturar audio del sistema, es necesario instalar la biblioteca sounddevice: "
                                "pip install sounddevice")
                return False

            # Enumerar dispositivos de audio
            dispositivos = sd.query_devices()

            # Buscar dispositivos de loopback (que capturan el audio del sistema)
            dispositivos_loopback = []
            for i, dispositivo in enumerate(dispositivos):
                nombre = dispositivo["name"].lower()
                if "loopback" in nombre or "stereo mix" in nombre or "mezcla estéreo" in nombre:
                    dispositivos_loopback.append((i, dispositivo["name"]))

            if not dispositivos_loopback:
                messagebox.showwarning("Advertencia", 
                                    "No se encontraron dispositivos de loopback (Stereo Mix). "
                                    "Es posible que necesite habilitarlo en la configuración de sonido de Windows.")
                return False

            # Configurar el dispositivo de captura como el primer dispositivo de loopback encontrado
            indice_dispositivo = dispositivos_loopback[0][0]

            # Actualizar la lista de dispositivos para asegurarnos de que contiene todos los dispositivos
            self.actualizar_dispositivos_audio()

            # Encontrar el índice en la lista de dispositivos
            try:
                idx_lista = self.dispositivos_nombres.index(dispositivos_loopback[0][1])
                self.dispositivos_combo.current(idx_lista)
                self.dispositivo_idx.set(indice_dispositivo)
            except ValueError:
                # Si no se encuentra en la lista, simplemente establecer el índice directamente
                self.dispositivo_idx.set(indice_dispositivo)

            messagebox.showinfo("Éxito", 
                            f"Se ha configurado el dispositivo de captura a: {dispositivos_loopback[0][1]}")
            return True

        except Exception as e:
            messagebox.showerror("Error", f"No se pudo configurar la captura de audio del sistema: {str(e)}")
            return False
        










    def actualizar_dispositivos_audio(self):
        """Actualiza la lista de dispositivos de audio disponibles"""
        dispositivos = []
        for i in range(self.audio.get_device_count()):
            try:
                info = self.audio.get_device_info_by_index(i)
                if info['maxInputChannels'] > 0:
                    nombre = info['name']
                    dispositivos.append((nombre, i))
            except:
                pass

        # Actualizar la lista de nombres de dispositivos
        self.dispositivos_nombres = [nombre for nombre, _ in dispositivos]

        # Actualizar el combobox
        self.dispositivos_combo.config(values=self.dispositivos_nombres)

        # Seleccionar el dispositivo actual si existe
        if dispositivos and self.dispositivo_idx.get() < len(dispositivos):
            self.dispositivos_combo.current(self.dispositivo_idx.get())
        elif dispositivos:
            self.dispositivos_combo.current(0)

    def detectar_y_configurar_audio_sistema(self):
        """Detecta el sistema operativo y configura la captura de audio del sistema"""
        sistema = platform.system()

        if sistema == "Windows":
            self.capturar_audio_sistema()
        elif sistema == "Darwin":  # macOS
            messagebox.showinfo("Configuración para macOS", 
                            "Para capturar audio del sistema en macOS:\n\n"
                            "1. Instale BlackHole desde: https://existential.audio/blackhole/\n"
                            "2. Configure BlackHole como dispositivo de salida\n"
                            "3. En Configuración de Audio MIDI, cree un dispositivo agregado\n"
                            "4. Seleccione BlackHole como dispositivo de entrada en esta aplicación")
        elif sistema == "Linux":
            messagebox.showinfo("Configuración para Linux", 
                            "Para capturar audio del sistema en Linux:\n\n"
                            "1. Ejecute: pactl load-module module-loopback\n"
                            "2. Seleccione el dispositivo 'Monitor of...' en la lista de dispositivos")
        else:
            messagebox.showwarning("Sistema no soportado",
                            f"La captura de audio del sistema no está soportada en {sistema}.")

    def iniciar_transcripcion(self):
        """Inicia la grabación y transcripción del audio."""
        print("Método iniciar_transcripcion() llamado")
        print(f"Modo: {self.modo_reconocimiento.get()}, Idioma: {self.idioma_var.get()}")

        if self.grabacion_activa:
            messagebox.showwarning("Advertencia", "La grabación ya está en curso.")
            return

        self.grabacion_activa = True
        self.frames = []
        self.estado_var.set("Grabando...")

        try:
            # Obtener el índice del dispositivo de audio seleccionado
            dispositivo_idx = self.dispositivos_combo.current()
            print(f"Usando dispositivo de índice: {dispositivo_idx}")

            self.stream = self.audio.open(
                format=pyaudio.paInt16,
                channels=1,
                rate=16000,
                input=True,
                input_device_index=dispositivo_idx,
                frames_per_buffer=1024,
            )

            def grabar_audio():
                print("Hilo de grabación iniciado")
                while self.grabacion_activa:
                    try:
                        data = self.stream.read(1024, exception_on_overflow=False)
                        self.frames.append(data)
                    except Exception as e:
                        print(f"Error durante la grabación: {e}")
                        # Añadir un pequeño tiempo de espera antes de reintentar
                        time.sleep(0.1)
                        continue  # Intentar de nuevo en lugar de romper el bucle

                print("Hilo de grabación detenido")

            self.hilo_grabacion = threading.Thread(target=grabar_audio)
            self.hilo_grabacion.daemon = True
            self.hilo_grabacion.start()

            # Iniciar contador de tiempo
            self.iniciar_contador_tiempo()

            # Deshabilitar el botón de iniciar y habilitar el de detener
            self.boton_iniciar.config(state=tk.DISABLED)
            self.boton_detener.config(state=tk.NORMAL)

            if self.modo_reconocimiento.get() == "online":
                threading.Thread(target=self.transcribir_online, daemon=True).start()
            elif self.modo_reconocimiento.get() == "offline":
                threading.Thread(target=self.transcribir_offline, daemon=True).start()

        except Exception as e:
            print(f"Error al iniciar la transcripción: {e}")
            self.estado_var.set("Error al iniciar")
            messagebox.showerror("Error", f"No se pudo iniciar la transcripción: {str(e)}")
            self.grabacion_activa = False
            if self.stream:
                try:
                    self.stream.stop_stream()
                    self.stream.close()
                except:
                    pass
            try:
                self.audio.terminate()
                self.audio = pyaudio.PyAudio()  # Reinicializar PyAudio
            except:
                pass
            self.boton_iniciar.config(state=tk.NORMAL)
            self.boton_detener.config(state=tk.DISABLED)

    def iniciar_contador_tiempo(self):
        """Inicia el contador de tiempo de grabación"""
        self.tiempo_inicio = datetime.now()
        self.actualizando_tiempo = True

        def actualizar_tiempo():
            if not self.actualizando_tiempo:
                return

            if not self.grabacion_activa:
                self.actualizando_tiempo = False
                return

            tiempo_transcurrido = datetime.now() - self.tiempo_inicio
            segundos_totales = int(tiempo_transcurrido.total_seconds())
            horas = segundos_totales // 3600
            minutos = (segundos_totales % 3600) // 60
            segundos = segundos_totales % 60

            tiempo_formateado = f"{horas:02d}:{minutos:02d}:{segundos:02d}"
            self.tiempo_var.set(tiempo_formateado)

            # Estimar tamaño del audio (16 bits, 16kHz, mono)
            kb_por_segundo = 16 * 16 * 1 / 8 / 1024  # kB por segundo
            tamano_estimado = kb_por_segundo * segundos_totales

            if tamano_estimado < 1024:
                self.tamano_var.set(f"{tamano_estimado:.2f} KB")
            else:
                self.tamano_var.set(f"{tamano_estimado / 1024:.2f} MB")

            # Programar la próxima actualización
            self.root.after(1000, actualizar_tiempo)

        # Iniciar la actualización
        actualizar_tiempo()

    def detener_transcripcion(self):
        """Detiene la grabación y (opcionalmente) guarda el audio."""
        print("Método detener_transcripcion() llamado")

        if not self.grabacion_activa:
            messagebox.showwarning("Advertencia", "No hay ninguna grabación en curso.")
            return

        self.grabacion_activa = False
        self.estado_var.set("Deteniendo...")

        try:
            # Esperar a que terminen los hilos
            tiempo_espera = 0
            while self.actualizando_tiempo and tiempo_espera < 30:
                time.sleep(0.1)
                tiempo_espera += 1

            if self.stream:
                self.stream.stop_stream()
                self.stream.close()
                self.stream = None

            # Asegurarse de reinicializar PyAudio correctamente
            self.audio.terminate()
            self.audio = pyaudio.PyAudio()

            # Guardar el audio capturado si hay frames
            if self.frames:
                # Procesar el audio completo para una transcripción final
                threading.Thread(target=self.procesar_audio_completo, daemon=True).start()

                # Guardar el audio en un archivo temporal
                ruta_audio = self.guardar_audio_temporal()
                if ruta_audio:
                    print(f"Audio guardado en: {ruta_audio}")

            # Habilitar el botón de iniciar y deshabilitar el de detener
            # Usar root.after para asegurar que se ejecuta en el hilo principal
            self.root.after(0, lambda: self.boton_iniciar.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.boton_detener.config(state=tk.DISABLED))

            self.estado_var.set("Listo")
            print("Grabación detenida correctamente")

        except Exception as e:
            print(f"Error al detener la transcripción: {e}")
            self.estado_var.set("Error al detener")
            messagebox.showerror("Error", f"No se pudo detener la transcripción: {str(e)}")

            # Asegurar que los botones estén en el estado correcto en caso de error
            self.root.after(0, lambda: self.boton_iniciar.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.boton_detener.config(state=tk.DISABLED))

    def guardar_audio_temporal(self):
        """Guarda el audio capturado en un archivo temporal (para depuración)."""
        if self.frames:
            nombre_archivo = f"audio_temporal_{datetime.now().strftime('%Y%m%d_%H%M%S')}.wav"
            ruta_archivo = os.path.join(os.path.dirname(os.path.abspath(__file__)), nombre_archivo)
            try:
                with wave.open(ruta_archivo, 'wb') as wf:
                    wf.setnchannels(1)  # Mono
                    wf.setsampwidth(self.audio.get_sample_size(pyaudio.paInt16))
                    wf.setframerate(16000)
                    wf.writeframes(b''.join(self.frames))
                print(f"Audio guardado en: {ruta_archivo}")  # Depuración
                return ruta_archivo
            except Exception as e:
                print(f"Error al guardar audio temporal: {e}")  # Depuración
                return None
        return None
    
    
    
    
    
    
    
    
    
    
    def transcribir_online(self):
        """Realiza la transcripción online usando Google Speech Recognition"""
        print("Método transcribir_online() llamado")

        if not self.grabacion_activa:
            return

        idioma = self.idioma_var.get()
        print(f"Iniciando transcripción online en idioma: {idioma}")

        buffer_frames = []
        contador = 0

        while self.grabacion_activa:
            try:
                # Esperar a tener suficientes frames para procesar
                while len(buffer_frames) < 16000 * self.duracion_segmento / 1024 and self.grabacion_activa:
                    if len(self.frames) > contador:
                        buffer_frames.append(self.frames[contador])
                        contador += 1
                    else:
                        time.sleep(0.1)

                if not self.grabacion_activa:
                    break

                # Crear archivo temporal para el segmento actual
                with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_file:
                    ruta_temp = temp_file.name

                # Guardar el segmento en el archivo temporal
                with wave.open(ruta_temp, 'wb') as wf:
                    wf.setnchannels(1)
                    wf.setsampwidth(self.audio.get_sample_size(pyaudio.paInt16))
                    wf.setframerate(16000)
                    wf.writeframes(b''.join(buffer_frames))

                # Transcribir el segmento
                with sr.AudioFile(ruta_temp) as fuente:
                    try:
                        audio_data = self.reconocedor.record(fuente)
                        texto = self.reconocedor.recognize_google(audio_data, language=idioma)

                        if texto.strip():
                            texto = self.aplicar_correcciones_post(texto, idioma)

                            # Calcular el tiempo del segmento
                            tiempo_seg = contador * 1024 / 16000
                            horas = int(tiempo_seg // 3600)
                            minutos = int((tiempo_seg % 3600) // 60)
                            segundos = int(tiempo_seg % 60)
                            tiempo_formateado = f"[{horas:02d}:{minutos:02d}:{segundos:02d}]"

                            # Detectar palabras clave en el texto
                            texto_con_etiquetas = self.detectar_palabras_clave(texto)

                            # Actualizar la transcripción con el tiempo
                            self.root.after(0, lambda t=tiempo_formateado, txt=texto_con_etiquetas: 
                                        self.actualizar_transcripcion_con_tiempo(t, txt))
                    except sr.UnknownValueError:
                        # No se detectó texto claro
                        pass
                    except Exception as e:
                        print(f"Error en transcripción online: {e}")

                # Eliminar archivo temporal
                try:
                    os.unlink(ruta_temp)
                except:
                    pass

                # Mantener un buffer deslizante con superposición
                frames_a_mantener = int((16000 * self.superposicion / 1024))
                buffer_frames = buffer_frames[-frames_a_mantener:] if frames_a_mantener > 0 else []

            except Exception as e:
                print(f"Error en el bucle de transcripción online: {e}")
                time.sleep(0.5)

    def actualizar_transcripcion_con_tiempo(self, tiempo, texto):
        """Actualiza el área de transcripción añadiendo una marca de tiempo"""
        if isinstance(texto, str):
            self.texto_transcripcion.insert(tk.END, f"{tiempo} {texto}\n\n")
        else:
            # Si texto es una lista de elementos (texto normal y con etiquetas)
            self.texto_transcripcion.insert(tk.END, tiempo + " ")
            for item in texto:
                if isinstance(item, tuple):
                    etiqueta, contenido = item
                    self.texto_transcripcion.insert(tk.END, contenido, etiqueta)
                else:
                    self.texto_transcripcion.insert(tk.END, item)
            self.texto_transcripcion.insert(tk.END, "\n\n")

        self.texto_transcripcion.see(tk.END)

    def actualizar_transcripcion(self, texto):
        """Actualiza el área de transcripción sin añadir marca de tiempo"""
        if isinstance(texto, str):
            self.texto_transcripcion.insert(tk.END, texto + "\n\n")
        else:
            # Si texto es una lista de elementos (texto normal y con etiquetas)
            for item in texto:
                if isinstance(item, tuple):
                    etiqueta, contenido = item
                    self.texto_transcripcion.insert(tk.END, contenido, etiqueta)
                else:
                    self.texto_transcripcion.insert(tk.END, item)
            self.texto_transcripcion.insert(tk.END, "\n\n")

        self.texto_transcripcion.see(tk.END)

    def transcribir_offline(self):
        """Realiza la transcripción offline usando Vosk"""
        print("Método transcribir_offline() llamado")

        if not vosk_disponible or not self.grabacion_activa:
            print("Vosk no disponible o grabación no activa")
            return

        idioma = self.idioma_var.get()
        idioma_base = idioma.split('-')[0]

        # Verificar si tenemos un modelo para este idioma
        if idioma_base not in self.modelos_vosk:
            self.estado_var.set(f"No hay modelo para {idioma_base}")
            print(f"No se encontró modelo Vosk para el idioma: {idioma_base}")
            self.root.after(0, lambda: messagebox.showwarning("Modelo no disponible", 
                                    f"No se encontró modelo para el idioma {idioma_base}. "
                                    "Se intentará usar el modo online."))
            self.modo_reconocimiento.set("online")
            self.transcribir_online()
            return

        try:
            # Cargar el modelo en este punto, ya que antes solo guardamos la ruta
            ruta_modelo = self.modelos_vosk[idioma_base]
            if not os.path.exists(ruta_modelo):
                raise FileNotFoundError(f"La ruta del modelo {ruta_modelo} no existe")

            print(f"Cargando modelo Vosk desde: {ruta_modelo}")
            model = Model(ruta_modelo)
            rec = KaldiRecognizer(model, 16000)

            buffer_size = 4000  # Tamaño del buffer para procesar
            contador = 0
            ultimo_texto = ""

            while self.grabacion_activa:
                # Procesar datos de audio en bloques
                if len(self.frames) > contador:
                    data = self.frames[contador]
                    contador += 1

                    if rec.AcceptWaveform(data):
                        result = json.loads(rec.Result())
                        if "text" in result and result["text"].strip():
                            texto = result["text"]

                            # Evitar repeticiones
                            if texto != ultimo_texto:
                                ultimo_texto = texto
                                texto = self.aplicar_correcciones_post(texto, idioma)

                                # Calcular el tiempo
                                tiempo_seg = contador * 1024 / 16000
                                horas = int(tiempo_seg // 3600)
                                minutos = int((tiempo_seg % 3600) // 60)
                                segundos = int(tiempo_seg % 60)
                                tiempo_formateado = f"[{horas:02d}:{minutos:02d}:{segundos:02d}]"

                                # Detectar palabras clave
                                texto_con_etiquetas = self.detectar_palabras_clave(texto)

                                # Actualizar la transcripción
                                self.root.after(0, lambda t=tiempo_formateado, txt=texto_con_etiquetas: 
                                            self.actualizar_transcripcion_con_tiempo(t, txt))
                else:
                    # Esperar a que haya más frames
                    time.sleep(0.1)

                    # Procesar lo que tenemos si llevamos tiempo sin procesar
                    if contador % 50 == 0:
                        result = json.loads(rec.PartialResult())
                        if "partial" in result and result["partial"].strip():
                            texto = result["partial"]
                            if texto != ultimo_texto and len(texto) > 10:  # Solo mostrar si es significativo
                                ultimo_texto = texto
                                texto = self.aplicar_correcciones_post(texto, idioma)

                                # Añadir indicador de parcial
                                texto = f"(parcial) {texto}"

                                # Calcular el tiempo
                                tiempo_seg = contador * 1024 / 16000
                                horas = int(tiempo_seg // 3600)
                                minutos = int((tiempo_seg % 3600) // 60)
                                segundos = int(tiempo_seg % 60)
                                tiempo_formateado = f"[{horas:02d}:{minutos:02d}:{segundos:02d}]"

                                # Detectar palabras clave
                                texto_con_etiquetas = self.detectar_palabras_clave(texto)

                                # Actualizar la transcripción
                                self.root.after(0, lambda t=tiempo_formateado, txt=texto_con_etiquetas: 
                                            self.actualizar_transcripcion_con_tiempo(t, txt))

        except Exception as e:
            print(f"Error en transcripción offline: {e}")
            self.estado_var.set(f"Error: {str(e)}")
            # Intentar usar el modo online como fallback
            if self.grabacion_activa:
                self.modo_reconocimiento.set("online")
                self.transcribir_online()
                
                
                
                
                
                
    def aplicar_correcciones_post(self, texto, idioma):
        """Aplica correcciones post-procesamiento al texto reconocido"""
        if not texto:
            return texto

        # Normalizar espacios
        texto = " ".join(texto.split())

        # Asegurar que la primera letra sea mayúscula
        if texto and len(texto) > 0:
            texto = texto[0].upper() + texto[1:]

        # Asegurar que termine con punto si no tiene otro signo de puntuación final
        if texto and not texto[-1] in ['.', '!', '?', ':', ';']:
            texto += '.'

        # Correcciones específicas por idioma
        if idioma.startswith('es'):
            # Ejemplo: Corregir "q" por "que"
            texto = texto.replace(" q ", " que ")
            texto = texto.replace("Q ", "Que ")

            # Otras correcciones específicas para español
            reemplazos_es = {
                " xq ": " porque ",
                " xa ": " para ",
                " tb ": " también ",
                " x ": " por ",
            }

            for original, reemplazo in reemplazos_es.items():
                texto = texto.replace(original, reemplazo)

        elif idioma.startswith('en'):
            # Correcciones específicas para inglés
            reemplazos_en = {
                " u ": " you ",
                " r ": " are ",
                " y ": " why ",
                " btw ": " by the way ",
            }

            for original, reemplazo in reemplazos_en.items():
                texto = texto.replace(original, reemplazo)

        return texto

    def aplicar_correcciones_contexto(self, texto, transcripciones_previas):
        """Aplica correcciones basadas en el contexto de transcripciones anteriores"""
        if not texto or not transcripciones_previas:
            return texto

        # Si hay superposición con la última transcripción, intentar fusionar
        ultima_transcripcion = transcripciones_previas[-1]

        # Buscar los últimos N caracteres de la transcripción anterior en la actual
        for n in range(min(30, len(ultima_transcripcion)), 5, -5):
            suffix = ultima_transcripcion[-n:]
            if texto.startswith(suffix):
                # Hay superposición, fusionar eliminando la parte duplicada
                return ultima_transcripcion + texto[n:]

        # Si no hay superposición pero la última transcripción no termina con
        # puntuación, considerar que es continuación
        if ultima_transcripcion and not ultima_transcripcion[-1] in ['.', '!', '?']:
            if not texto[0].isupper():  # Si la nueva transcripción no comienza con mayúscula
                return ultima_transcripcion + " " + texto

        return texto

    def procesar_audio_completo(self):
        """Procesa el audio capturado completo para transcripción final"""
        if not self.frames:
            return

        self.estado_var.set("Procesando transcripción final...")
        print("Procesando el audio completo para transcripción final")

        try:
            # Crear archivo temporal para el audio completo
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_file:
                ruta_temp = temp_file.name

            # Guardar el audio capturado en el archivo temporal
            with wave.open(ruta_temp, 'wb') as wf:
                wf.setnchannels(1)
                wf.setsampwidth(self.audio.get_sample_size(pyaudio.paInt16))
                wf.setframerate(16000)
                wf.writeframes(b''.join(self.frames))

            # Añadir delimitador en la transcripción
            self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, "\n\n----- Transcripción Final -----\n\n"))

            idioma = self.idioma_var.get()
            modo = self.modo_reconocimiento.get()

            # Procesar el audio según el modo seleccionado
            if modo == "online":
                try:
                    with sr.AudioFile(ruta_temp) as fuente:
                        self.reconocedor.adjust_for_ambient_noise(fuente, duration=min(1.0, 0.5))
                        audio_data = self.reconocedor.record(fuente)
                        texto = self.reconocedor.recognize_google(audio_data, language=idioma)

                        if texto.strip():
                            texto = self.aplicar_correcciones_post(texto, idioma)
                            texto_con_etiquetas = self.detectar_palabras_clave(texto)
                            self.root.after(0, lambda: self.actualizar_transcripcion(texto_con_etiquetas))

                except sr.UnknownValueError:
                    self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, "No se pudo reconocer claramente el audio completo.\n"))
                except Exception as e:
                    print(f"Error en transcripción final online: {e}")
                    self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, f"Error en la transcripción final: {str(e)}\n"))

            elif modo == "offline" and vosk_disponible:
                try:
                    idioma_base = idioma.split('-')[0]

                    if idioma_base in self.modelos_vosk:
                        ruta_modelo = self.modelos_vosk[idioma_base]
                        model = Model(ruta_modelo)
                        rec = KaldiRecognizer(model, 16000)

                        with open(ruta_temp, "rb") as wf:
                            wf.read(44)  # Saltar cabecera WAV
                            buffer_size = 4000

                            while True:
                                data = wf.read(buffer_size)
                                if len(data) == 0:
                                    break

                                if rec.AcceptWaveform(data):
                                    result = json.loads(rec.Result())
                                    if "text" in result and result["text"].strip():
                                        texto = result["text"]
                                        texto = self.aplicar_correcciones_post(texto, idioma)
                                        texto_con_etiquetas = self.detectar_palabras_clave(texto)
                                        self.root.after(0, lambda t=texto_con_etiquetas: self.actualizar_transcripcion(t))

                            # Procesar el último fragmento
                            result = json.loads(rec.FinalResult())
                            if "text" in result and result["text"].strip():
                                texto = result["text"]
                                texto = self.aplicar_correcciones_post(texto, idioma)
                                texto_con_etiquetas = self.detectar_palabras_clave(texto)
                                self.root.after(0, lambda t=texto_con_etiquetas: self.actualizar_transcripcion(t))
                    else:
                        self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, f"No se encontró modelo para el idioma {idioma_base}.\n"))

                except Exception as e:
                    print(f"Error en transcripción final offline: {e}")
                    self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, f"Error en la transcripción offline: {str(e)}\n"))

            # Eliminar archivo temporal
            try:
                os.unlink(ruta_temp)
            except:
                pass

            self.estado_var.set("Transcripción completada")

        except Exception as e:
            print(f"Error al procesar audio completo: {e}")
            self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, f"Error al procesar audio: {str(e)}\n"))
            self.estado_var.set("Error en transcripción")

    def cargar_modelos_vosk(self):
        """Carga los modelos Vosk disponibles para reconocimiento offline"""
        if not vosk_disponible:
            return

        dir_modelos = os.path.join(os.path.dirname(os.path.abspath(__file__)), "modelos")
        if not os.path.exists(dir_modelos):
            try:
                os.makedirs(dir_modelos)
            except Exception as e:
                print(f"Error creando directorio de modelos: {e}")

        try:
            for item in os.listdir(dir_modelos):
                ruta_item = os.path.join(dir_modelos, item)
                if os.path.isdir(ruta_item) and item.startswith("vosk-model"):
                    # Determinar el idioma del modelo
                    if "es" in item.lower():
                        self.modelos_vosk["es"] = ruta_item  # Solo guardamos la ruta, no el modelo cargado
                    elif "en" in item.lower():
                        self.modelos_vosk["en"] = ruta_item
                    else:
                        # Modelo genérico o de otro idioma
                        codigo_idioma = item.split("-")[-1].lower() if "-" in item else "generic"
                        self.modelos_vosk[codigo_idioma] = ruta_item

            print(f"Encontrados {len(self.modelos_vosk)} modelos Vosk: {', '.join(self.modelos_vosk.keys())}")
        except Exception as e:
            print(f"Error al cargar modelos Vosk: {e}")

    def detectar_palabras_clave(self, texto):
        """Detecta palabras clave en el texto y las marca para resaltado"""
        if not self.palabras_clave or not texto:
            return texto

        try:
            # Contar palabras clave encontradas
            palabras_encontradas = 0

            # Si no hay palabras clave definidas, devolver el texto sin modificar
            if not self.palabras_clave:
                return texto

            # Convertir texto a minúsculas para comparación insensible a mayúsculas
            texto_lower = texto.lower()

            # Lista para almacenar el texto con etiquetas
            resultado = []

            # Posición actual en el texto
            pos_actual = 0

            # Buscar cada palabra clave en el texto
            for palabra in self.palabras_clave:
                palabra_lower = palabra.lower()
                inicio = 0

                # Buscar todas las ocurrencias de la palabra clave
                while True:
                    indice = texto_lower.find(palabra_lower, inicio)
                    if indice == -1:
                        break

                    # Verificar si es una palabra completa
                    es_palabra_completa = True
                    if indice > 0 and texto_lower[indice-1].isalnum():
                        es_palabra_completa = False
                    fin_palabra = indice + len(palabra_lower)
                    if fin_palabra < len(texto_lower) and texto_lower[fin_palabra].isalnum():
                        es_palabra_completa = False

                    if es_palabra_completa:
                        # Añadir el texto anterior a la palabra clave
                        if indice > pos_actual:
                            resultado.append(texto[pos_actual:indice])

                        # Añadir la palabra clave con su etiqueta
                        resultado.append(("palabra_clave", texto[indice:indice+len(palabra_lower)]))

                        pos_actual = indice + len(palabra_lower)
                        palabras_encontradas += 1

                    inicio = indice + 1

            # Añadir el resto del texto después de la última palabra clave
            if pos_actual < len(texto):
                resultado.append(texto[pos_actual:])

            # Actualizar contador de palabras clave
            if palabras_encontradas > 0:
                self.root.after(0, lambda: self.palabras_detectadas_var.set(str(palabras_encontradas)))

            return resultado if resultado else texto

        except Exception as e:
            print(f"Error al detectar palabras clave: {e}")
            return texto
        
        
        
        
        
        
        
        
        
        
        
        
        
        
        
    def buscar_modelos_vosk(self):
        """Busca modelos de Vosk en el sistema"""
        if not vosk_disponible:
            messagebox.showwarning("Vosk no instalado", "Primero debe instalar Vosk: pip install vosk")
            return

        dir_modelos = os.path.join(os.path.dirname(os.path.abspath(__file__)), "modelos")
        if not os.path.exists(dir_modelos):
            try:
                os.makedirs(dir_modelos)
            except Exception as e:
                print(f"Error creando directorio de modelos: {e}")

        modelos_encontrados = []
        try:
            for item in os.listdir(dir_modelos):
                ruta_item = os.path.join(dir_modelos, item)
                if os.path.isdir(ruta_item) and item.startswith("vosk-model"):
                    modelos_encontrados.append(item)
        except Exception as e:
            print(f"Error al buscar modelos Vosk: {e}")

        if modelos_encontrados:
            mensaje = "Modelos encontrados:\n"
            for modelo in modelos_encontrados:
                mensaje += f"- {modelo}\n"
            messagebox.showinfo("Modelos encontrados", mensaje)
            self.cargar_modelos_vosk()
        else:
            messagebox.showinfo(
                "Seleccionar modelos", 
                "No se encontraron modelos en la carpeta 'modelos/'. "
                "Seleccione la carpeta donde ha descargado los modelos de Vosk."
            )

            carpeta_modelos = filedialog.askdirectory(title="Seleccionar carpeta con modelos Vosk")
            if carpeta_modelos:
                modelos_en_carpeta = []
                try:
                    for item in os.listdir(carpeta_modelos):
                        ruta_item = os.path.join(carpeta_modelos, item)
                        if os.path.isdir(ruta_item) and item.startswith("vosk-model"):
                            modelos_en_carpeta.append(item)
                except Exception as e:
                    print(f"Error al verificar carpeta de modelos: {e}")

                if modelos_en_carpeta:
                    respuesta = messagebox.askyesno(
                        "Modelos encontrados", 
                        f"Se encontraron {len(modelos_en_carpeta)} modelos. "
                        "¿Desea copiarlos a la carpeta 'modelos/' del programa?"
                    )
                    if respuesta:
                        for modelo in modelos_en_carpeta:
                            origen = os.path.join(carpeta_modelos, modelo)
                            destino = os.path.join(dir_modelos, modelo)
                            self.estado_var.set(f"Copiando modelo {modelo}...")
                            try:
                                shutil.copytree(origen, destino)
                                print(f"Modelo copiado: {modelo}")
                            except Exception as e:
                                print(f"Error al copiar modelo {modelo}: {e}")

                        self.estado_var.set("Modelos copiados")
                        messagebox.showinfo("Éxito", "Modelos copiados correctamente")
                        self.cargar_modelos_vosk()
                else:
                    messagebox.showwarning(
                        "No se encontraron modelos", 
                        "No se encontraron modelos de Vosk en la carpeta seleccionada."
                    )

    def guardar_palabras_clave(self):
        """Guarda las palabras clave en un archivo"""
        texto_palabras = self.texto_palabras_clave.get("1.0", tk.END).strip()
        palabras = [p.strip().lower() for p in texto_palabras.split('\n') if p.strip()]

        palabras = list(set(palabras))

        self.palabras_clave = set(palabras)

        archivo_palabras_clave = os.path.join(os.path.dirname(os.path.abspath(__file__)), "palabras_clave.txt")

        try:
            with open(archivo_palabras_clave, 'w', encoding='utf-8') as f:
                for palabra in sorted(palabras):
                    f.write(palabra + '\n')

            messagebox.showinfo("Éxito", f"Se guardaron {len(palabras)} palabras clave")
            print(f"Guardadas {len(palabras)} palabras clave")

        except Exception as e:
            print(f"Error al guardar palabras clave: {e}")
            messagebox.showerror("Error", f"No se pudieron guardar las palabras clave: {str(e)}")

    def anadir_palabra_clave(self):
        """Añade una nueva palabra clave"""
        nueva_palabra = simpledialog.askstring(
            "Nueva Palabra Clave", 
            "Ingrese la nueva palabra clave corporativa:"
        )

        if nueva_palabra and nueva_palabra.strip():
            palabra = nueva_palabra.strip().lower()

            texto_actual = self.texto_palabras_clave.get("1.0", tk.END).strip()
            if palabra not in texto_actual:
                if texto_actual:
                    self.texto_palabras_clave.insert(tk.END, f"\n{palabra}")
                else:
                    self.texto_palabras_clave.insert(tk.END, palabra)

                print(f"Añadida palabra clave: {palabra}")

    def eliminar_palabra_clave(self):
        """Elimina la palabra clave seleccionada"""
        try:
            seleccion = self.texto_palabras_clave.tag_ranges(tk.SEL)
            if seleccion:
                texto_seleccionado = self.texto_palabras_clave.get(seleccion[0], seleccion[1])
                self.texto_palabras_clave.delete(seleccion[0], seleccion[1])
                print(f"Eliminada palabra clave: {texto_seleccionado.strip()}")
            else:
                messagebox.showinfo("Información", "Seleccione una palabra clave para eliminar")
        except:
            messagebox.showinfo("Información", "Seleccione una palabra clave para eliminar")

    def guardar_configuracion(self):
        """Guarda la configuración actual"""
        try:
            config = {
                "duracion_segmento": self.duracion_segmento,
                "superposicion": self.superposicion,
                "energy_threshold": self.reconocedor.energy_threshold,
                "pause_threshold": self.reconocedor.pause_threshold,
                "dynamic_energy_threshold": self.reconocedor.dynamic_energy_threshold,
                "cancelacion_ruido": self.cancelacion_ruido_var.get(),
                "formato_salida": self.formato_salida.get(),
                "dispositivo_idx": self.dispositivo_idx.get(),
                "modo_privado": self.modo_privado.get(),
                "usar_encriptacion": self.usar_encriptacion.get(),
                "modo_reconocimiento": self.modo_reconocimiento.get(),
                "idioma": self.idioma_var.get()
            }

            archivo_config = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")

            with open(archivo_config, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=4)

            print("Configuración guardada")
            messagebox.showinfo("Éxito", "Configuración guardada correctamente")

        except Exception as e:
            print(f"Error al guardar configuración: {e}")
            messagebox.showerror("Error", f"No se pudo guardar la configuración: {str(e)}")

    def establecer_clave_encriptacion(self):
        """Solicita al usuario ingresar una clave de encriptación"""
        if not encriptacion_disponible:
            messagebox.showwarning("No disponible", "La encriptación no está disponible")
            return False

        clave = simpledialog.askstring(
            "Clave de Encriptación", 
            "Ingrese una clave para encriptar/desencriptar (mínimo 8 caracteres):",
            show='*'
        )

        if not clave:
            return False

        if len(clave) < 8:
            messagebox.showwarning("Clave débil", "La clave debe tener al menos 8 caracteres")
            return False

        confirmacion = simpledialog.askstring(
            "Confirmar Clave", 
            "Confirme la clave de encriptación:",
            show='*'
        )

        if clave != confirmacion:
            messagebox.showerror("Error", "Las claves no coinciden")
            return False

        self.clave_encriptacion = clave
        self.salt = os.urandom(16)

        print("Nueva clave de encriptación establecida")
        return True

    def encriptar_transcripcion(self):
        """Encripta la transcripción actual"""
        if not encriptacion_disponible:
            messagebox.showwarning("No disponible", "La encriptación no está disponible. Instale la biblioteca cryptography.")
            return

        texto = self.texto_transcripcion.get("1.0", tk.END).strip()
        if not texto:
            messagebox.showwarning("Advertencia", "No hay texto para encriptar")
            return

        if not self.clave_encriptacion:
            resultado = self.establecer_clave_encriptacion()
            if not resultado:
                return

        try:
            kdf = PBKDF2HMAC(
                algorithm=hashes.SHA256(),
                length=32,
                salt=self.salt,
                iterations=100000,
            )

            key = base64.urlsafe_b64encode(kdf.derive(self.clave_encriptacion.encode()))
            cipher_suite = Fernet(key)

            contenido_encriptado = cipher_suite.encrypt(texto.encode('utf-8'))

            self.texto_transcripcion.delete("1.0", tk.END)
            self.texto_transcripcion.insert(tk.END, "[CONTENIDO ENCRIPTADO]\n\n")
            self.texto_transcripcion.insert(tk.END, "Este texto ha sido encriptado y solo puede ser leído con la clave correcta.\n")
            self.texto_transcripcion.insert(tk.END, "Use 'Guardar Transcripción' para guardar el archivo encriptado.\n")

            self.contenido_encriptado = contenido_encriptado

            messagebox.showinfo("Éxito", "Transcripción encriptada. Utilice 'Guardar Transcripción' para guardarla.")
            print("Transcripción encriptada en memoria")

        except Exception as e:
            print(f"Error al encriptar transcripción: {e}")
            messagebox.showerror("Error", f"Error al encriptar: {str(e)}")

    def desencriptar_archivo(self):
        """Permite al usuario seleccionar y desencriptar un archivo"""
        if not encriptacion_disponible:
            messagebox.showwarning("No disponible", "La funcionalidad de desencriptación no está disponible")
            return

        archivo = filedialog.askopenfilename(
            title="Seleccionar Archivo Encriptado",
            filetypes=[
                ("Archivos Encriptados", "*.enc"),
                ("Todos los Archivos", "*.*")
            ]
        )

        if not archivo:
            return

        try:
            with open(archivo, 'rb') as f:
                contenido = f.read()

            # Verificar formato
            if contenido[:6] == b'TENC01':  # Formato nuevo
                header = contenido[:6]
                salt = contenido[6:22]
                contenido_encriptado = contenido[22:]
            else:  # Formato antiguo o incompatible
                salt = contenido[:16]
                contenido_encriptado = contenido[16:]

            # Solicitar clave si no existe
            if not self.clave_encriptacion:
                clave = simpledialog.askstring(
                    "Clave de Desencriptación", 
                    "Ingrese la clave para desencriptar el archivo:",
                    show='*'
                )

                if not clave:
                    return

                self.clave_encriptacion = clave

            # Derivar clave
            kdf = PBKDF2HMAC(
                algorithm=hashes.SHA256(),
                length=32,
                salt=salt,
                iterations=100000,
            )

            key = base64.urlsafe_b64encode(kdf.derive(self.clave_encriptacion.encode()))
            cipher_suite = Fernet(key)

            # Intentar desencriptar
            try:
                texto_desencriptado = cipher_suite.decrypt(contenido_encriptado).decode('utf-8')

                # Mostrar el contenido desencriptado
                self.texto_transcripcion.delete("1.0", tk.END)
                self.texto_transcripcion.insert(tk.END, texto_desencriptado)

                self.estado_var.set("Archivo desencriptado")
                self.logger.info(f"Archivo desencriptado: {os.path.basename(archivo)}")

                messagebox.showinfo("Éxito", "Archivo desencriptado correctamente")

            except Exception as e:
                # Probablemente clave incorrecta
                messagebox.showerror("Error", "No se pudo desencriptar el archivo. La clave podría ser incorrecta.")
                self.clave_encriptacion = None  # Resetear la clave para el próximo intento
                self.logger.warning(f"Intento fallido de desencriptar archivo: {os.path.basename(archivo)}")

        except Exception as e:
            print(f"Error al desencriptar archivo: {e}")
            self.logger.error(f"Error al procesar archivo encriptado: {str(e)}")
            messagebox.showerror("Error", f"Error al procesar el archivo: {str(e)}")
            
            
            
            
            
            
            
            
    def guardar_transcripcion(self):
        """Guarda la transcripción actual en un archivo de texto"""
        texto = self.texto_transcripcion.get("1.0", tk.END).strip()

        if not texto:
            messagebox.showwarning("Advertencia", "No hay texto para guardar")
            return

        fecha_hora = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_defecto = f"transcripcion_{fecha_hora}"

        formato_elegido = self.formato_salida.get()
        extension = f".{formato_elegido}"

        if formato_elegido == "enc" and encriptacion_disponible:
            if not self.clave_encriptacion:
                resultado = self.establecer_clave_encriptacion()
                if not resultado:
                    return

        archivo = filedialog.asksaveasfilename(
            title="Guardar Transcripción",
            defaultextension=extension,
            initialfile=nombre_defecto,
            filetypes=self.obtener_tipos_archivo_guardar()
        )

        if archivo:
            try:
                extension = os.path.splitext(archivo)[1].lower()

                if extension == '.docx':
                    try:
                        import docx

                        doc = docx.Document()

                        doc.add_heading('Transcripción de Audio', 0)

                        doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
                        doc.add_paragraph(f'Idioma: {self.idioma_var.get()}')

                        doc.add_paragraph(texto)

                        doc.save(archivo)

                    except ImportError:
                        messagebox.showinfo("Información", 
                                        "Librería python-docx no instalada. Guardando como texto plano.")
                        with open(archivo, 'w', encoding='utf-8') as f:
                            f.write(texto)

                elif extension == '.srt':
                    subtitulos = []

                    lineas = texto.split('\n')
                    for i, linea in enumerate(lineas):
                        if linea.startswith('[') and ']' in linea:
                            partes = linea.split(']', 1)
                            if len(partes) == 2:
                                tiempo_str = partes[0][1:]
                                texto_subtitulo = partes[1].strip()

                                h, m, s = tiempo_str.split(':')
                                tiempo_inicio = f"{h}:{m}:{s},000"

                                h_fin, m_fin, s_fin = int(h), int(m), int(s) + 5
                                if s_fin >= 60:
                                    s_fin -= 60
                                    m_fin += 1
                                if m_fin >= 60:
                                    m_fin -= 60
                                    h_fin += 1

                                tiempo_fin = f"{h_fin:02d}:{m_fin:02d}:{s_fin:02d},000"

                                subtitulos.append({
                                    'num': len(subtitulos) + 1,
                                    'inicio': tiempo_inicio,
                                    'fin': tiempo_fin,
                                    'texto': texto_subtitulo
                                })

                    with open(archivo, 'w', encoding='utf-8') as f:
                        for sub in subtitulos:
                            f.write(f"{sub['num']}\n")
                            f.write(f"{sub['inicio']} --> {sub['fin']}\n")
                            f.write(f"{sub['texto']}\n\n")

                elif extension == '.enc' and encriptacion_disponible:
                    self.guardar_archivo_encriptado(archivo, texto)

                else:
                    with open(archivo, 'w', encoding='utf-8') as f:
                        f.write(texto)

                print(f"Transcripción guardada: {os.path.basename(archivo)}")
                messagebox.showinfo("Éxito", f"Transcripción guardada en: {archivo}")

            except Exception as e:
                print(f"Error al guardar archivo: {e}")
                messagebox.showerror("Error", f"Error al guardar el archivo: {str(e)}")

    def obtener_tipos_archivo_guardar(self):
        """Devuelve los tipos de archivo disponibles para guardar"""
        tipos = [
            ("Archivos de Texto", "*.txt"),
            ("Todos los Archivos", "*.*")
        ]

        if docx_disponible:
            tipos.insert(1, ("Documentos Word", "*.docx"))

        tipos.insert(1, ("Documentos SRT", "*.srt"))

        if encriptacion_disponible:
            tipos.insert(1, ("Archivos Encriptados", "*.enc"))

        return tipos

    def guardar_archivo_encriptado(self, ruta_archivo, contenido):
        """Guarda un archivo con encriptación segura"""
        if not encriptacion_disponible:
            raise Exception("Biblioteca de encriptación no disponible")

        if not self.clave_encriptacion:
            # Solicitar clave si no existe
            resultado = self.establecer_clave_encriptacion()
            if not resultado:
                raise Exception("No se ha establecido clave de encriptación")

        try:
            # Generar un nuevo salt para cada archivo (más seguro)
            salt = os.urandom(16)

            # Derivar clave con PBKDF2
            kdf = PBKDF2HMAC(
                algorithm=hashes.SHA256(),
                length=32,
                salt=salt,
                iterations=100000,  # Aumentar iteraciones para mayor seguridad
            )

            key = base64.urlsafe_b64encode(kdf.derive(self.clave_encriptacion.encode()))
            cipher_suite = Fernet(key)

            # Encriptar el contenido
            contenido_encriptado = cipher_suite.encrypt(contenido.encode('utf-8'))

            # Guardar salt y contenido encriptado
            with open(ruta_archivo, 'wb') as f:
                # Guardar formato y versión (para compatibilidad futura)
                f.write(b'TENC01')  # Identificador y versión
                # Guardar salt
                f.write(salt)
                # Guardar contenido encriptado
                f.write(contenido_encriptado)

            # Registrar la acción en el log
            self.logger.info(f"Archivo encriptado guardado: {os.path.basename(ruta_archivo)}")
            print(f"Archivo guardado con encriptación: {os.path.basename(ruta_archivo)}")

            return True

        except Exception as e:
            print(f"Error en encriptación: {e}")
            self.logger.error(f"Error al encriptar archivo: {str(e)}")
            raise Exception(f"Error al encriptar: {str(e)}")

    def toggle_grabacion(self):
        """Inicia o detiene la grabación de audio"""
        if not hasattr(self, 'grabando'):
            self.grabando = False

        if not self.grabando:
            # Iniciar grabación
            self.grabando = True
            self.estado_var.set("Grabando audio...")
            self.boton_grabar.config(text="Detener Grabación")
            # Iniciar la transcripción
            self.iniciar_transcripcion()
            print("Grabación iniciada")
        else:
            # Detener grabación
            self.grabando = False
            self.estado_var.set("Audio detenido")
            self.boton_grabar.config(text="Iniciar Grabación")
            # Detener la transcripción
            self.detener_transcripcion()
            print("Grabación detenida")

    def cargar_archivo(self):
        """Permite al usuario seleccionar un archivo de audio para transcribir"""
        # Determinar los formatos soportados
        filetypes = [
            ("Archivos de audio", "*.wav;*.mp3;*.ogg;*.flac"),
            ("Archivos WAV", "*.wav"),
            ("Archivos MP3", "*.mp3"),
            ("Archivos OGG", "*.ogg"),
            ("Archivos FLAC", "*.flac"),
            ("Todos los Archivos", "*.*")
        ]

        # Permitir al usuario seleccionar un archivo
        archivo = filedialog.askopenfilename(
            title="Seleccionar Archivo de Audio",
            filetypes=filetypes
        )

        if archivo:
            print(f"Archivo seleccionado: {archivo}")
            # Crear ventana de opciones para la transcripción
            ventana_opciones = tk.Toplevel(self.root)
            ventana_opciones.title("Opciones de Transcripción")
            ventana_opciones.geometry("400x300")
            ventana_opciones.grab_set()

            frame_opciones = ttk.Frame(ventana_opciones, padding=10)
            frame_opciones.pack(fill=tk.BOTH, expand=True)

            ttk.Label(
                frame_opciones,
                text=f"Archivo: {os.path.basename(archivo)}",
                font=("Segoe UI", 10, "bold")
            ).pack(pady=(0, 10))

            # Modo de transcripción
            ttk.Label(frame_opciones, text="Modo de reconocimiento:").pack(anchor=tk.W)

            modo_archivo = tk.StringVar(value=self.modo_reconocimiento.get())

            ttk.Radiobutton(
                frame_opciones,
                text="Online (Google)",
                value="online",
                variable=modo_archivo
            ).pack(anchor=tk.W, padx=20)

            ttk.Radiobutton(
                frame_opciones,
                text="Offline (Vosk)",
                value="offline",
                variable=modo_archivo,
                state=tk.NORMAL if vosk_disponible else tk.DISABLED
            ).pack(anchor=tk.W, padx=20)

            # Idioma
            ttk.Label(frame_opciones, text="Idioma:").pack(anchor=tk.W, pady=(10, 0))

            idioma_archivo = tk.StringVar(value=self.idioma_var.get())

            idiomas = [
                ("Español (ES)", "es-ES"),
                ("Inglés (US)", "en-US"),
                ("Inglés (UK)", "en-GB"),
                ("Español (MX)", "es-MX")
            ]

            for texto, valor in idiomas:
                ttk.Radiobutton(
                    frame_opciones,
                    text=texto,
                    value=valor,
                    variable=idioma_archivo
                ).pack(anchor=tk.W, padx=20)

            # Opciones adicionales
            ttk.Label(frame_opciones, text="Opciones adicionales:").pack(anchor=tk.W, pady=(10, 0))

            cancelar_ruido = tk.BooleanVar(value=self.cancelacion_ruido_var.get())

            ttk.Checkbutton(
                frame_opciones,
                text="Aplicar cancelación de ruido",
                variable=cancelar_ruido,
                state=tk.NORMAL if pydub_disponible else tk.DISABLED
            ).pack(anchor=tk.W, padx=20)

            # Botones
            frame_botones = ttk.Frame(frame_opciones)
            frame_botones.pack(fill=tk.X, pady=(20, 0))

            def iniciar_transcripcion_archivo():
                modo = modo_archivo.get()
                idioma = idioma_archivo.get()

                ventana_opciones.destroy()

                # Iniciar la transcripción del archivo
                threading.Thread(
                    target=self.transcribir_archivo_grande,
                    args=(archivo, modo),
                    daemon=True
                ).start()

            ttk.Button(
                frame_botones,
                text="Transcribir",
                command=iniciar_transcripcion_archivo
            ).pack(side=tk.LEFT)

            ttk.Button(
                frame_botones,
                text="Cancelar",
                command=ventana_opciones.destroy
            ).pack(side=tk.RIGHT)

    def limpiar_transcripcion(self):
        """Limpia el área de transcripción"""
        self.texto_transcripcion.delete("1.0", tk.END)
        self.palabras_detectadas_var.set("0")
        print("Transcripción limpiada")

    def actualizar_estado_idioma(self):
        """Actualiza los indicadores de estado según el idioma seleccionado"""
        idioma = self.idioma_var.get()
        modo = self.modo_reconocimiento.get()

        if modo == "offline":
            idioma_base = idioma.split('-')[0]
            if idioma_base not in self.modelos_vosk:
                messagebox.showwarning(
                    "Modelo no disponible", 
                    f"No se encontró modelo para el idioma {idioma_base} en modo offline. "
                    "Se usará el modo online para este idioma."
                )
                self.modo_reconocimiento.set("online")
                self.actualizar_estado_modo()

    def cambiar_modo_reconocimiento(self):
        """Cambia entre modos de reconocimiento online/offline"""
        modo = self.modo_reconocimiento.get()

        if modo == "offline":
            if not vosk_disponible:
                messagebox.showwarning(
                    "Vosk no instalado", 
                    "El reconocimiento offline requiere la biblioteca Vosk. Instálela con: pip install vosk"
                )
                self.modo_reconocimiento.set("online")
                return

            if not self.modelos_vosk:
                messagebox.showwarning(
                    "Modelos no disponibles", 
                    "No se encontraron modelos de idioma para Vosk. Descargue los modelos e instálelos en la carpeta 'modelos/'."
                )

        self.actualizar_estado_modo()
        print(f"Modo de reconocimiento cambiado a: {modo}")

    def actualizar_estado_modo(self):
        """Actualiza los indicadores de estado según el modo seleccionado"""
        modo = self.modo_reconocimiento.get()

        if modo == "online":
            self.icono_privacidad_var.set("🔓 Modo estándar (envía audio a servidores externos)")
            self.label_privacidad.configure(foreground="#e74c3c")
            self.barra_estado.configure(text="Listo para transcribir | Modo: Online (Google API)")
        else:
            self.icono_privacidad_var.set("🔒 Modo privado (procesamiento local)")
            self.label_privacidad.configure(foreground="#27ae60")
            self.barra_estado.configure(text="Listo para transcribir | Modo: Offline (Vosk)")
            
            
            
            
            
            
            
    def transcribir_archivo_grande(self, archivo_audio, modo):
        """Transcribe un archivo de audio grande dividiéndolo en segmentos"""
        try:
            # Verificar la extensión del archivo
            extension = os.path.splitext(archivo_audio)[1].lower()

            # Si no es un archivo WAV, intentar convertirlo
            if extension != '.wav':
                if pydub_disponible:
                    self.root.after(0, lambda: self.estado_var.set(f"Convirtiendo archivo {extension} a WAV..."))

                    # Crear un archivo temporal
                    temp_wav = tempfile.NamedTemporaryFile(suffix='.wav', delete=False).name

                    # Convertir el archivo al formato WAV
                    try:
                        sound = AudioSegment.from_file(archivo_audio)
                        sound.export(temp_wav, format="wav")
                        archivo_audio = temp_wav  # Usar el archivo WAV temporal
                        self.root.after(0, lambda: self.estado_var.set("Archivo convertido, procesando..."))
                    except Exception as e:
                        print(f"Error al convertir archivo: {e}")
                        self.root.after(0, lambda: self.estado_var.set(f"Error al convertir archivo: {str(e)}"))
                        messagebox.showerror("Error", f"Error al convertir archivo: {str(e)}")
                        return
                else:
                    messagebox.showerror("Error", 
                                f"No se puede procesar archivos {extension}. Instale pydub: pip install pydub")
                    return

            # Continuar con el procesamiento del archivo WAV
            with wave.open(archivo_audio, 'rb') as wf:
                canales = wf.getnchannels()
                tasa = wf.getframerate()

                frames = wf.getnframes()
                rate = wf.getframerate()
                duracion = frames / float(rate)

                info_audio = (f"Información del archivo:\n"
                            f"- Duración: {int(duracion // 60)}:{int(duracion % 60):02d}\n"
                            f"- Canales: {canales}\n"
                            f"- Tasa de muestreo: {tasa} Hz\n"
                            f"- Tamaño: {os.path.getsize(archivo_audio) / (1024*1024):.2f} MB\n\n")

                self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, info_audio))

                print(f"Transcribiendo archivo: {os.path.basename(archivo_audio)}, "
                    f"Duración: {int(duracion // 60)}:{int(duracion % 60):02d}, "
                    f"Modo: {modo}")

                if duracion > 60:
                    self.root.after(0, lambda: self.estado_var.set("Procesando archivo grande por segmentos..."))
                    self.procesar_archivo_por_segmentos(archivo_audio, duracion, modo)
                else:
                    self.transcribir_segmento_archivo(archivo_audio, modo)

            # Limpiar el archivo temporal si se creó uno
            if extension != '.wav' and pydub_disponible and 'temp_wav' in locals():
                try:
                    os.unlink(temp_wav)
                except:
                    pass

        except Exception as e:
            print(f"Error al procesar archivo: {e}")
            self.root.after(0, lambda: self.estado_var.set(f"Error al procesar archivo: {str(e)}"))
            messagebox.showerror("Error", f"Error al procesar archivo: {str(e)}")

    def transcribir_segmento_archivo(self, archivo_audio, modo):
        """Transcribe un archivo de audio completo de una sola vez"""
        try:
            idioma = self.idioma_var.get()
            
            if modo == "online":
                with sr.AudioFile(archivo_audio) as fuente:
                    self.reconocedor.adjust_for_ambient_noise(fuente, duration=min(1.0, 0.5))
                    audio_data = self.reconocedor.record(fuente)
                    
                    try:
                        texto = self.reconocedor.recognize_google(audio_data, language=idioma)
                        
                        if texto.strip():
                            texto = self.aplicar_correcciones_post(texto, idioma)
                            texto_con_etiquetas = self.detectar_palabras_clave(texto)
                            self.root.after(0, lambda: self.actualizar_transcripcion(texto_con_etiquetas))
                            
                    except sr.UnknownValueError:
                        self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, 
                            "No se pudo reconocer el audio.\n"))
                    except Exception as e:
                        print(f"Error en transcripción: {e}")
                        self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, 
                            f"Error en la transcripción: {str(e)}\n"))
                        
            elif modo == "offline" and vosk_disponible:
                idioma_base = idioma.split('-')[0]
                
                if idioma_base in self.modelos_vosk:
                    model = Model(self.modelos_vosk[idioma_base])
                    rec = KaldiRecognizer(model, self.tasa_muestreo)
                    
                    with open(archivo_audio, 'rb') as wf:
                        wf.read(44)  # Saltar cabecera WAV
                        
                        while True:
                            data = wf.read(4000)
                            if len(data) == 0:
                                break
                                
                            rec.AcceptWaveform(data)
                            
                        result = json.loads(rec.FinalResult())
                        if "text" in result and result["text"].strip():
                            texto = result["text"]
                            texto = self.aplicar_correcciones_post(texto, idioma)
                            texto_con_etiquetas = self.detectar_palabras_clave(texto)
                            self.root.after(0, lambda: self.actualizar_transcripcion(texto_con_etiquetas))
                            
            self.root.after(0, lambda: self.estado_var.set("Transcripción completada"))
            
        except Exception as e:
            print(f"Error al transcribir segmento: {e}")
            self.root.after(0, lambda: self.estado_var.set(f"Error: {str(e)}"))

    def procesar_archivo_por_segmentos(self, archivo_audio, duracion_total, modo):
        """Procesa un archivo de audio grande por segmentos con solapamiento"""
        try:
            idioma = self.idioma_var.get()
            idioma_base = idioma.split('-')[0]

            tamano_segmento = 15
            solapamiento = 2

            num_segmentos = math.ceil((duracion_total - solapamiento) / (tamano_segmento - solapamiento))

            self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, 
                            f"Archivo dividido en {num_segmentos} segmentos con solapamiento...\n\n"))

            transcripciones_previas = []

            reconocedor_offline = None
            if modo == "offline" and vosk_disponible and idioma_base in self.modelos_vosk:
                modelo = Model(self.modelos_vosk[idioma_base])
                reconocedor_offline = KaldiRecognizer(modelo, self.tasa_muestreo)

            # Procesar usando archivos temporales para cada segmento
            for i in range(num_segmentos):
                progreso = int((i+1) * 100 / num_segmentos)
                self.root.after(0, lambda i=i, p=progreso: 
                            self.estado_var.set(f"Procesando segmento {i+1}/{num_segmentos} ({p}%)..."))

                offset = i * (tamano_segmento - solapamiento)
                duracion = min(tamano_segmento, duracion_total - offset)

                if duracion <= 0:
                    break

                try:
                    # Crear archivo temporal para el segmento
                    with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_file:
                        temp_path = temp_file.name

                    # Extraer el segmento de audio
                    with wave.open(archivo_audio, 'rb') as wf_in:
                        params = wf_in.getparams()
                        framerate = params.framerate
                        
                        # Calcular posiciones en frames
                        start_frame = int(offset * framerate)
                        num_frames = int(duracion * framerate)
                        
                        # Posicionarse en el inicio del segmento
                        wf_in.setpos(start_frame)
                        
                        # Leer los frames del segmento
                        frames_data = wf_in.readframes(num_frames)
                        
                        # Guardar el segmento en archivo temporal
                        with wave.open(temp_path, 'wb') as wf_out:
                            wf_out.setparams(params)
                            wf_out.writeframes(frames_data)

                    # Transcribir el segmento
                    if modo == "online":
                        with sr.AudioFile(temp_path) as fuente:
                            self.reconocedor.adjust_for_ambient_noise(fuente, duration=0.5)
                            audio_segmento = self.reconocedor.record(fuente)
                            texto = self.reconocedor.recognize_google(audio_segmento, language=idioma)
                    else:
                        if reconocedor_offline:
                            with open(temp_path, 'rb') as wf:
                                wf.read(44)  # Saltar cabecera WAV
                                audio_data = wf.read()
                                
                                reconocedor_offline.AcceptWaveform(audio_data)
                                resultado = json.loads(reconocedor_offline.FinalResult())
                                texto = resultado.get('text', '')

                    # Eliminar archivo temporal
                    try:
                        os.unlink(temp_path)
                    except:
                        pass

                    if texto and texto.strip():
                        if transcripciones_previas:
                            texto = self.aplicar_correcciones_contexto(texto, transcripciones_previas)

                        texto = self.aplicar_correcciones_post(texto, idioma)

                        transcripciones_previas.append(texto)
                        if len(transcripciones_previas) > 5:
                            transcripciones_previas = transcripciones_previas[-5:]

                        horas_inicio = int(offset // 3600)
                        min_inicio = int((offset % 3600) // 60)
                        seg_inicio = int(offset % 60)

                        tiempo_formateado = f"[{horas_inicio:02d}:{min_inicio:02d}:{seg_inicio:02d}]"

                        texto_con_etiquetas = self.detectar_palabras_clave(texto)

                        self.root.after(0, lambda tiempo=tiempo_formateado, texto=texto_con_etiquetas: 
                                    self.actualizar_transcripcion_con_tiempo(tiempo, texto))

                except sr.UnknownValueError:
                    pass

                except Exception as e:
                    print(f"Error en segmento {i+1}: {e}")

            self.root.after(0, lambda: self.estado_var.set("Transcripción completada"))
            self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, 
                        "\n--- Fin de la transcripción ---\n"))

            print(f"Transcripción de archivo completada: {os.path.basename(archivo_audio)}")

        except Exception as e:
            print(f"Error al procesar archivo por segmentos: {e}")
            self.root.after(0, lambda: self.estado_var.set(f"Error: {str(e)}"))
            
            
            
            
            
            
            
            
            
    def crear_interfaz(self):
        """Crea la interfaz gráfica del usuario"""
        # Configurar estilo
        self.configurar_estilo()

        # Frame principal con pestañas
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Pestaña principal de transcripción
        tab_principal = ttk.Frame(notebook)
        notebook.add(tab_principal, text="Transcripción")

        # Pestaña de configuración
        tab_config = ttk.Frame(notebook)
        notebook.add(tab_config, text="Configuración")

        # Pestaña de privacidad
        tab_privacidad = ttk.Frame(notebook)
        notebook.add(tab_privacidad, text="Privacidad")

        # Pestaña de ayuda
        tab_ayuda = ttk.Frame(notebook)
        notebook.add(tab_ayuda, text="Ayuda")

        # Configurar las pestañas
        self.configurar_tab_principal(tab_principal)
        self.configurar_tab_config(tab_config)
        self.configurar_tab_privacidad(tab_privacidad)
        self.configurar_tab_ayuda(tab_ayuda)

        # Barra de estado en la parte inferior
        self.barra_estado = ttk.Label(
            self.root, 
            text="Listo para transcribir | Modo: Online", 
            relief=tk.SUNKEN, 
            anchor=tk.W
        )
        self.barra_estado.pack(fill=tk.X, side=tk.BOTTOM, padx=10, pady=(5, 10))

    def configurar_estilo(self):
        """Configura el estilo de la aplicación"""
        style = ttk.Style()

        style.configure("TButton", padding=5, font=('Segoe UI', 10))
        style.configure("TLabel", font=('Segoe UI', 10))
        style.configure("TFrame", background="#f5f5f5")
        style.configure("TLabelframe", background="#f5f5f5")
        style.configure("TLabelframe.Label", font=('Segoe UI', 10, 'bold'))

        style.configure("Accent.TButton", font=('Segoe UI', 10, 'bold'))
        if os.name == 'nt':  # Windows
            style.map('Accent.TButton',
                background=[('active', '!disabled', '#3498db')],
                foreground=[('active', '!disabled', 'white')])

        style.configure("Privacy.TButton", font=('Segoe UI', 10))
        style.map('Privacy.TButton',
            background=[('active', '!disabled', '#27ae60')],
            foreground=[('active', '!disabled', 'white')])

    def configurar_tab_principal(self, parent):
        """Configura la pestaña principal de transcripción"""
        main_frame = ttk.Frame(parent)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        opciones_frame = ttk.LabelFrame(main_frame, text="Opciones de Transcripción", padding=10)
        opciones_frame.pack(fill=tk.X, pady=(0, 10))

        opciones_frame.columnconfigure(0, weight=1)
        opciones_frame.columnconfigure(1, weight=1)

        idioma_frame = ttk.Frame(opciones_frame)
        idioma_frame.grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)

        ttk.Label(idioma_frame, text="Idioma:").pack(side=tk.LEFT, padx=(0, 10))

        idiomas = [
            ("Español (ES)", "es-ES"),
            ("Inglés (US)", "en-US"),
            ("Inglés (UK)", "en-GB"),
            ("Español (MX)", "es-MX")
        ]

        for i, (texto, valor) in enumerate(idiomas):
            ttk.Radiobutton(
                idioma_frame, 
                text=texto, 
                value=valor, 
                variable=self.idioma_var,
                command=self.actualizar_estado_idioma
            ).pack(side=tk.LEFT, padx=(0, 10))

        modo_frame = ttk.Frame(opciones_frame)
        modo_frame.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)

        ttk.Label(modo_frame, text="Modo:").pack(side=tk.LEFT, padx=(0, 10))

        modos = [
            ("Online (Google)", "online"),
            ("Offline (Vosk)", "offline")
        ]

        for texto, valor in modos:
            rb = ttk.Radiobutton(
                modo_frame,
                text=texto,
                value=valor,
                variable=self.modo_reconocimiento,
                command=self.cambiar_modo_reconocimiento
            )
            rb.pack(side=tk.LEFT, padx=(0, 10))

            if valor == "offline" and not vosk_disponible:
                rb.configure(state=tk.DISABLED)

        privacidad_frame = ttk.Frame(opciones_frame)
        privacidad_frame.grid(row=1, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)

        self.icono_privacidad_var = tk.StringVar(value="🔓 Modo estándar (envía audio a servidores externos)")
        self.label_privacidad = ttk.Label(
            privacidad_frame, 
            textvariable=self.icono_privacidad_var,
            foreground="#e74c3c"
        )
        self.label_privacidad.pack(side=tk.LEFT, padx=(0, 10))

        stats_frame = ttk.LabelFrame(main_frame, text="Información de Grabación", padding=10)
        stats_frame.pack(fill=tk.X, pady=(0, 10))

        for i in range(3):
            stats_frame.columnconfigure(i, weight=1)

        ttk.Label(stats_frame, text="Tiempo de grabación:").grid(row=0, column=0, sticky=tk.W)
        ttk.Label(stats_frame, textvariable=self.tiempo_var, font=("Segoe UI", 10, "bold")).grid(row=0, column=1, sticky=tk.W)

        ttk.Label(stats_frame, text="Tamaño estimado:").grid(row=1, column=0, sticky=tk.W)
        ttk.Label(stats_frame, textvariable=self.tamano_var).grid(row=1, column=1, sticky=tk.W)

        ttk.Label(stats_frame, text="Estado:").grid(row=0, column=2, sticky=tk.W)
        ttk.Label(stats_frame, textvariable=self.estado_var, font=("Segoe UI", 10, "bold")).grid(row=0, column=3, sticky=tk.W)

        ttk.Label(stats_frame, text="Palabras clave:").grid(row=1, column=2, sticky=tk.W)
        ttk.Label(stats_frame, textvariable=self.palabras_detectadas_var, foreground="#27ae60").grid(row=1, column=3, sticky=tk.W)

        ttk.Label(main_frame, text="Transcripción en tiempo real:", font=("Segoe UI", 11, "bold")).pack(anchor=tk.W, pady=(5, 5))

        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

        self.texto_transcripcion = tk.Text(text_frame, wrap=tk.WORD, height=15, font=("Segoe UI", 11))
        self.texto_transcripcion.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.texto_transcripcion.tag_configure("palabra_clave", foreground="#27ae60", font=("Segoe UI", 11, "bold"))
        self.texto_transcripcion.tag_configure("advertencia", foreground="#e74c3c", font=("Segoe UI", 11, "bold"))

        scrollbar = ttk.Scrollbar(text_frame, command=self.texto_transcripcion.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.texto_transcripcion.config(yscrollcommand=scrollbar.set)

        action_frame = ttk.Frame(main_frame)
        action_frame.pack(fill=tk.X, pady=10)

        self.boton_grabar = ttk.Button(
            action_frame,
            text="Iniciar Grabación", 
            command=self.toggle_grabacion,
            style="Accent.TButton"
        )
        self.boton_grabar.pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(
            action_frame,
            text="Cargar Archivo", 
            command=self.cargar_archivo
        ).pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(
            action_frame,
            text="Guardar Transcripción", 
            command=self.guardar_transcripcion
        ).pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(
            action_frame,
            text="Limpiar",
            command=self.limpiar_transcripcion
        ).pack(side=tk.LEFT)

        self.boton_encriptar = ttk.Button(
            action_frame,
            text="Encriptar",
            command=self.encriptar_transcripcion,
            style="Privacy.TButton"
        )
        self.boton_encriptar.pack(side=tk.RIGHT)
        if not encriptacion_disponible:
            self.boton_encriptar.configure(state=tk.DISABLED)

        # Botones de iniciar y detener
        self.boton_iniciar = ttk.Button(
            action_frame,
            text="Iniciar Transcripción",
            command=self.iniciar_transcripcion,
        )
        self.boton_iniciar.pack(side=tk.LEFT, padx=(10, 10))

        self.boton_detener = ttk.Button(
            action_frame,
            text="Detener Transcripción",
            command=self.detener_transcripcion,
            state=tk.DISABLED  # Deshabilitado al inicio
        )
        self.boton_detener.pack(side=tk.LEFT)

        ttk.Button(
            action_frame,
            text="Capturar Audio Sistema",
            command=self.detectar_y_configurar_audio_sistema,
            state=tk.NORMAL if sounddevice_disponible else tk.DISABLED
        ).pack(side=tk.LEFT, padx=(10, 0))

    def cambiar_modo_privado(self):
        """Activa o desactiva el modo privado"""
        if self.modo_privado.get():
            if not vosk_disponible or not self.modelos_vosk:
                respuesta = messagebox.askyesno(
                    "Componentes no disponibles", 
                    "El modo privado requiere Vosk para reconocimiento offline, pero no está disponible. "
                    "¿Desea continuar con el modo privado limitado (solo encriptación)?"
                )

                if not respuesta:
                    self.modo_privado.set(False)
                    return

            if vosk_disponible and self.modelos_vosk:
                self.modo_reconocimiento.set("offline")
                self.actualizar_estado_modo()

            if encriptacion_disponible:
                self.usar_encriptacion.set(True)
                self.actualizar_estado_encriptacion()

                if not self.clave_encriptacion:
                    self.establecer_clave_encriptacion()

            print("Modo privado activado")
        else:
            print("Modo privado desactivado")

    def actualizar_estado_encriptacion(self):
        """Actualiza el estado de encriptación"""
        if self.usar_encriptacion.get() and encriptacion_disponible:
            self.formato_salida.set("enc")
        else:
            self.formato_salida.set("txt")
            
            
            
            
    def configurar_tab_config(self, parent):
        """Configura la pestaña de configuración"""
        config_frame = ttk.Frame(parent)
        config_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        audio_frame = ttk.LabelFrame(config_frame, text="Configuración de Audio", padding=10)
        audio_frame.pack(fill=tk.X, pady=(0, 10))

        for i in range(4):
            audio_frame.columnconfigure(i, weight=1)

        ttk.Label(audio_frame, text="Duración segmento (seg):").grid(row=0, column=0, sticky=tk.W, padx=5)
        self.duracion_var = tk.StringVar(value=str(self.duracion_segmento))
        duracion_spin = ttk.Spinbox(
            audio_frame,
            from_=1,
            to=10,
            width=5,
            textvariable=self.duracion_var,
            command=lambda: setattr(self, 'duracion_segmento', int(self.duracion_var.get()))
        )
        duracion_spin.grid(row=0, column=1, sticky=tk.W, padx=5)

        ttk.Label(audio_frame, text="Superposición (seg):").grid(row=0, column=2, sticky=tk.W, padx=5)
        self.superposicion_var = tk.StringVar(value=str(self.superposicion))
        superposicion_spin = ttk.Spinbox(
            audio_frame,
            from_=0.5,
            to=3,
            increment=0.5,
            width=5,
            textvariable=self.superposicion_var,
            command=lambda: setattr(self, 'superposicion', float(self.superposicion_var.get()))
        )
        superposicion_spin.grid(row=0, column=3, sticky=tk.W, padx=5)

        ttk.Label(audio_frame, text="Umbral de energía:").grid(row=1, column=0, sticky=tk.W, padx=5)
        self.energia_var = tk.StringVar(value=str(self.reconocedor.energy_threshold))
        energia_spin = ttk.Spinbox(
            audio_frame,
            from_=50,
            to=1000,
            increment=50,
            width=5,
            textvariable=self.energia_var,
            command=lambda: setattr(self.reconocedor, 'energy_threshold', int(self.energia_var.get()))
        )
        energia_spin.grid(row=1, column=1, sticky=tk.W, padx=5)

        ttk.Label(audio_frame, text="Umbral de pausa (seg):").grid(row=1, column=2, sticky=tk.W, padx=5)
        self.pausa_var = tk.StringVar(value=str(self.reconocedor.pause_threshold))
        pausa_spin = ttk.Spinbox(
            audio_frame,
            from_=0.3,
            to=2.0,
            increment=0.1,
            width=5,
            textvariable=self.pausa_var,
            command=lambda: setattr(self.reconocedor, 'pause_threshold', float(self.pausa_var.get()))
        )
        pausa_spin.grid(row=1, column=3, sticky=tk.W, padx=5)

        opciones_frame = ttk.Frame(audio_frame)
        opciones_frame.grid(row=2, column=0, columnspan=4, sticky=tk.W, pady=5)

        self.ajuste_dinamico_var = tk.BooleanVar(value=self.reconocedor.dynamic_energy_threshold)
        ttk.Checkbutton(
            opciones_frame,
            text="Ajuste dinámico",
            variable=self.ajuste_dinamico_var,
            command=lambda: setattr(self.reconocedor, 'dynamic_energy_threshold', self.ajuste_dinamico_var.get())
        ).pack(side=tk.LEFT, padx=(0, 15))

        ttk.Checkbutton(
            opciones_frame,
            text="Cancelación de ruido",
            variable=self.cancelacion_ruido_var,
            state=tk.NORMAL if pydub_disponible else tk.DISABLED
        ).pack(side=tk.LEFT, padx=(0, 15))

        dispositivos_frame = ttk.LabelFrame(config_frame, text="Dispositivos de Audio", padding=10)
        dispositivos_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(dispositivos_frame, text="Dispositivo de entrada:").pack(anchor=tk.W, padx=5, pady=5)

        # Obtener lista de dispositivos de entrada
        dispositivos = []
        for i in range(self.audio.get_device_count()):
            try:
                info = self.audio.get_device_info_by_index(i)
                if info['maxInputChannels'] > 0:
                    nombre = info['name']
                    dispositivos.append((nombre, i))
            except:
                pass

        # Variable para almacenar los nombres de dispositivos
        self.dispositivos_nombres = [nombre for nombre, _ in dispositivos]

        # Combobox para seleccionar dispositivo
        self.dispositivos_combo = ttk.Combobox(
            dispositivos_frame,
            values=self.dispositivos_nombres,
            state="readonly",
            width=40
        )
        self.dispositivos_combo.pack(anchor=tk.W, padx=20, pady=5)

        # Seleccionar el dispositivo actual si existe
        if dispositivos and self.dispositivo_idx.get() < len(dispositivos):
            self.dispositivos_combo.current(self.dispositivo_idx.get())
        elif dispositivos:
            self.dispositivos_combo.current(0)

        # Información sobre dispositivos
        if not dispositivos:
            ttk.Label(
                dispositivos_frame,
                text="No se detectaron dispositivos de entrada",
                foreground="#e74c3c"
            ).pack(anchor=tk.W, padx=20, pady=5)
        else:
            ttk.Label(
                dispositivos_frame,
                text=f"Dispositivos detectados: {len(dispositivos)}",
                foreground="#27ae60"
            ).pack(anchor=tk.W, padx=20, pady=5)

        formato_frame = ttk.LabelFrame(config_frame, text="Formato de Salida", padding=10)
        formato_frame.pack(fill=tk.X, pady=(0, 10))

        formatos = [
            ("Texto plano (.txt)", "txt"),
            ("Documento Word (.docx)", "docx", docx_disponible),
            ("Subtítulos (.srt)", "srt"),
            ("Texto encriptado (.enc)", "enc", encriptacion_disponible)
        ]

        for texto, valor, *condicion in formatos:
            estado = tk.NORMAL if not condicion or condicion[0] else tk.DISABLED

            ttk.Radiobutton(
                formato_frame,
                text=texto,
                value=valor,
                variable=self.formato_salida,
                state=estado
            ).pack(anchor=tk.W, padx=20, pady=2)

        # Botones de acción
        botones_frame = ttk.Frame(config_frame)
        botones_frame.pack(fill=tk.X, pady=10)

        ttk.Button(
            botones_frame,
            text="Restaurar Valores Predeterminados",
            command=self.restaurar_config_predeterminada
        ).pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(
            botones_frame,
            text="Guardar Configuración",
            command=self.guardar_configuracion
        ).pack(side=tk.LEFT)
        
        
        
        
        
        
    def configurar_tab_privacidad(self, parent):
        """Configura la pestaña de privacidad"""
        privacidad_frame = ttk.Frame(parent)
        privacidad_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        modo_privado_frame = ttk.LabelFrame(privacidad_frame, text="Modo de Privacidad", padding=10)
        modo_privado_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Checkbutton(
            modo_privado_frame,
            text="Activar modo privado (usar reconocimiento offline y encriptación)",
            variable=self.modo_privado,
            command=self.cambiar_modo_privado
        ).pack(anchor=tk.W, padx=5, pady=5)

        # Información adicional sobre el modo privado
        ttk.Label(
            modo_privado_frame,
            text="El modo privado utiliza Vosk para procesamiento local sin enviar audio a servidores externos.\nSe recomienda para información sensible o confidencial.",
            wraplength=700
        ).pack(anchor=tk.W, padx=5, pady=5)

        estado_vosk = "Instalado y listo" if vosk_disponible and self.modelos_vosk else "No instalado"
        if vosk_disponible and not self.modelos_vosk:
            estado_vosk = "Instalado pero faltan modelos de idioma"

        ttk.Label(
            modo_privado_frame,
            text=f"Estado de Vosk (reconocimiento offline): {estado_vosk}",
            foreground="#27ae60" if (vosk_disponible and self.modelos_vosk) else "#e74c3c"
        ).pack(anchor=tk.W, padx=5, pady=5)

        if not vosk_disponible or not self.modelos_vosk:
            instrucciones_frame = ttk.Frame(modo_privado_frame)
            instrucciones_frame.pack(fill=tk.X, padx=5, pady=5)

            ttk.Label(
                instrucciones_frame,
                text="Instalación de Vosk para reconocimiento offline:",
                font=("Segoe UI", 10, "bold")
            ).pack(anchor=tk.W)

            instrucciones = (
                "1. Instalar Vosk: pip install vosk\n"
                "2. Descargar modelos desde https://alphacephei.com/vosk/models\n"
                "   - Para español: vosk-model-small-es\n"
                "   - Para inglés: vosk-model-small-en\n"
                "3. Descomprimir los modelos en la carpeta 'modelos/' del directorio del programa"
            )

            ttk.Label(
                instrucciones_frame,
                text=instrucciones,
                wraplength=700
            ).pack(anchor=tk.W, padx=20)

            ttk.Button(
                instrucciones_frame,
                text="Buscar modelos",
                command=self.buscar_modelos_vosk
            ).pack(anchor=tk.W, padx=20, pady=5)

        encriptacion_frame = ttk.LabelFrame(privacidad_frame, text="Encriptación", padding=10)
        encriptacion_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Checkbutton(
            encriptacion_frame,
            text="Activar encriptación de archivos guardados",
            variable=self.usar_encriptacion,
            state=tk.NORMAL if encriptacion_disponible else tk.DISABLED,
            command=self.actualizar_estado_encriptacion
        ).pack(anchor=tk.W, padx=5, pady=5)

        ttk.Label(
            encriptacion_frame,
            text="La encriptación protege los archivos de transcripción para que sólo puedan ser leídos con la clave correcta.",
            wraplength=700
        ).pack(anchor=tk.W, padx=5, pady=5)

        estado_encriptacion = "Disponible" if encriptacion_disponible else "No disponible"
        ttk.Label(
            encriptacion_frame,
            text=f"Estado de encriptación: {estado_encriptacion}",
            foreground="#27ae60" if encriptacion_disponible else "#e74c3c"
        ).pack(anchor=tk.W, padx=5, pady=5)

        if not encriptacion_disponible:
            ttk.Label(
                encriptacion_frame,
                text="Para habilitar la encriptación, instale la biblioteca cryptography:\npip install cryptography",
                wraplength=700
            ).pack(anchor=tk.W, padx=20, pady=5)

        botones_enc_frame = ttk.Frame(encriptacion_frame)
        botones_enc_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(
            botones_enc_frame,
            text="Establecer Clave",
            command=self.establecer_clave_encriptacion,
            state=tk.NORMAL if encriptacion_disponible else tk.DISABLED
        ).pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(
            botones_enc_frame,
            text="Desencriptar Archivo",
            command=self.desencriptar_archivo,
            state=tk.NORMAL if encriptacion_disponible else tk.DISABLED
        ).pack(side=tk.LEFT)

        keywords_frame = ttk.LabelFrame(privacidad_frame, text="Palabras Clave Corporativas", padding=10)
        keywords_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(
            keywords_frame,
            text="El transcriptor detectará estas palabras en las transcripciones y las resaltará para facilitar la identificación de información sensible.",
            wraplength=700
        ).pack(anchor=tk.W, padx=5, pady=5)

        palabras_actuales_frame = ttk.Frame(keywords_frame)
        palabras_actuales_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(
            palabras_actuales_frame,
            text="Palabras clave actuales:",
            font=("Segoe UI", 10, "bold")
        ).pack(anchor=tk.W)

        self.texto_palabras_clave = tk.Text(palabras_actuales_frame, height=5, width=50, font=("Segoe UI", 10))
        self.texto_palabras_clave.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(20, 5), pady=5)

        for palabra in sorted(self.palabras_clave):
            self.texto_palabras_clave.insert(tk.END, palabra + "\n")

        scrollbar_palabras = ttk.Scrollbar(palabras_actuales_frame, command=self.texto_palabras_clave.yview)
        scrollbar_palabras.pack(side=tk.RIGHT, fill=tk.Y)
        self.texto_palabras_clave.config(yscrollcommand=scrollbar_palabras.set)

        botones_keywords_frame = ttk.Frame(keywords_frame)
        botones_keywords_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(
            botones_keywords_frame,
            text="Guardar Palabras Clave",
            command=self.guardar_palabras_clave
        ).pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(
            botones_keywords_frame,
            text="Añadir Palabra Clave",
            command=self.anadir_palabra_clave
        ).pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(
            botones_keywords_frame,
            text="Eliminar Seleccionada",
            command=self.eliminar_palabra_clave
        ).pack(side=tk.LEFT)

        log_frame = ttk.LabelFrame(privacidad_frame, text="Registro de Actividad", padding=10)
        log_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(
            log_frame,
            text="El registro de actividad permite auditar el uso del transcriptor para fines de seguridad.\nRegistra información sobre sesiones de grabación, archivos procesados y acciones realizadas.",
            wraplength=700
        ).pack(anchor=tk.W, padx=5, pady=5)

        ttk.Button(
            log_frame,
            text="Ver Registros",
            command=self.ver_registros
        ).pack(anchor=tk.W, padx=5, pady=5)
        
        
        
        
        
        
        
    def configurar_tab_ayuda(self, parent):
        """Configura la pestaña de ayuda"""
        ayuda_frame = ttk.Frame(parent)
        ayuda_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        ttk.Label(
            ayuda_frame,
            text="Ayuda del Transcriptor Multilingüe",
            font=("Segoe UI", 14, "bold")
        ).pack(anchor=tk.W, pady=(0, 10))

        texto_ayuda = (
            "Bienvenido al Transcriptor Multilingüe con Privacidad Mejorada\n\n"
            "Este programa permite transcribir audio en español e inglés con opciones avanzadas de privacidad "
            "para proteger información confidencial o sensible. A continuación encontrará información sobre "
            "cómo utilizar las principales funciones:\n\n"

            "MODOS DE RECONOCIMIENTO\n"
            "------------------------\n"
            "1. Modo Online (Google): Utiliza la API de Google para reconocimiento de voz. Ofrece alta precisión "
            "pero envía el audio a servidores externos.\n\n"
            "2. Modo Offline (Vosk): Procesa el audio localmente sin enviar datos a servidores externos. "
            "Recomendado para información confidencial o sensible.\n\n"

            "PRIVACIDAD Y SEGURIDAD\n"
            "---------------------\n"
            "- Modo Privado: Activa automáticamente el reconocimiento offline y la encriptación.\n"
            "- Encriptación: Protege los archivos guardados con una clave personal.\n"
            "- Detección de Palabras Clave: Resalta términos corporativos o sensibles en las transcripciones.\n"
            "- Registro de Actividad: Mantiene un historial de uso para auditoría de seguridad.\n\n"

            "CONSEJOS PARA MEJOR RECONOCIMIENTO\n"
            "--------------------------------\n"
            "- Utilice un micrófono de buena calidad en un ambiente con poco ruido.\n"
            "- Hable claramente y a un ritmo normal.\n"
            "- Ajuste el umbral de energía según el nivel de ruido ambiental.\n"
            "- Para archivos grandes, utilice la opción de cancelación de ruido.\n"
            "- En modo offline, los modelos pequeños son más rápidos pero menos precisos.\n\n"

            "FORMATOS SOPORTADOS\n"
            "------------------\n"
            "- Entrada: WAV, MP3, OGG, FLAC (requiere pydub)\n"
            "- Salida: TXT, DOCX (requiere python-docx), SRT, ENC (encriptado)\n\n"

            "RESOLUCIÓN DE PROBLEMAS\n"
            "----------------------\n"
            "- Si el reconocimiento online falla, verifique su conexión a internet.\n"
            "- Si el reconocimiento offline no funciona, verifique que los modelos estén instalados correctamente.\n"
            "- Para problemas con la encriptación, reinstale la biblioteca cryptography.\n"
            "- Los archivos de registro se encuentran en la carpeta 'logs/' para diagnóstico.\n\n"

            "Para más información o soporte, contacte al administrador del sistema."
        )

        texto_ayuda_widget = tk.Text(ayuda_frame, wrap=tk.WORD, height=25, font=("Segoe UI", 10))
        texto_ayuda_widget.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        texto_ayuda_widget.insert(tk.END, texto_ayuda)
        texto_ayuda_widget.config(state=tk.DISABLED)

        scrollbar_ayuda = ttk.Scrollbar(texto_ayuda_widget, command=texto_ayuda_widget.yview)
        scrollbar_ayuda.pack(side=tk.RIGHT, fill=tk.Y)
        texto_ayuda_widget.config(yscrollcommand=scrollbar_ayuda.set)

        botones_ayuda_frame = ttk.Frame(ayuda_frame)
        botones_ayuda_frame.pack(fill=tk.X, pady=10)

        ttk.Button(
            botones_ayuda_frame,
            text="Verificar Instalación",
            command=self.verificar_instalacion
        ).pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(
            botones_ayuda_frame,
            text="Acerca de",
            command=self.mostrar_acerca_de
        ).pack(side=tk.LEFT)

    def restaurar_config_predeterminada(self):
        """Restaura la configuración predeterminada"""
        confirmar = messagebox.askyesno(
            "Confirmar", 
            "¿Está seguro de restaurar la configuración predeterminada?\nSe perderán todas las personalizaciones."
        )

        if not confirmar:
            return

        self.duracion_segmento = 3
        self.duracion_var.set(str(self.duracion_segmento))

        self.superposicion = 1
        self.superposicion_var.set(str(self.superposicion))

        self.reconocedor.energy_threshold = 300
        self.energia_var.set(str(self.reconocedor.energy_threshold))

        self.reconocedor.pause_threshold = 0.6
        self.pausa_var.set(str(self.reconocedor.pause_threshold))

        self.reconocedor.dynamic_energy_threshold = True
        self.ajuste_dinamico_var.set(self.reconocedor.dynamic_energy_threshold)

        self.cancelacion_ruido_var.set(True and pydub_disponible)

        self.formato_salida.set("txt")

        self.dispositivo_idx.set(0)

        self.modo_privado.set(False)
        self.usar_encriptacion.set(False)

        self.modo_reconocimiento.set("online")
        self.actualizar_estado_modo()

        print("Configuración restaurada a valores predeterminados")
        messagebox.showinfo("Éxito", "Configuración restaurada a valores predeterminados")

    def verificar_instalacion(self):
        """Verifica la instalación y muestra información de diagnóstico"""
        info_sistema = {
            "Python": f"{sys.version.split()[0]}",
            "OS": f"{os.name}",
            "Plataforma": f"{os.uname().sysname if hasattr(os, 'uname') else 'Windows'}"
        }

        componentes = {
            "SpeechRecognition": "Instalado" if 'sr' in globals() else "No instalado",
            "PyAudio": "Instalado" if 'pyaudio' in globals() else "No instalado",
            "Vosk": "Instalado" if vosk_disponible else "No instalado",
            "Cryptography": "Instalado" if encriptacion_disponible else "No instalado",
            "pydub": "Instalado" if pydub_disponible else "No instalado",
            "python-docx": "Instalado" if docx_disponible else "No instalado"
        }

        info_modelos = "No hay modelos de Vosk instalados"
        if self.modelos_vosk:
            info_modelos = f"Modelos instalados: {', '.join(self.modelos_vosk.keys())}"

        ventana_info = tk.Toplevel(self.root)
        ventana_info.title("Verificación de Instalación")
        ventana_info.geometry("600x500")
        ventana_info.minsize(600, 400)

        frame_info = ttk.Frame(ventana_info, padding=10)
        frame_info.pack(fill=tk.BOTH, expand=True)

        ttk.Label(
            frame_info,
            text="Diagnóstico de Instalación",
            font=("Segoe UI", 14, "bold")
        ).pack(anchor=tk.W, pady=(0, 10))

        ttk.Label(
            frame_info,
            text="Información del Sistema:",
            font=("Segoe UI", 11, "bold")
        ).pack(anchor=tk.W, pady=(10, 5))

        for clave, valor in info_sistema.items():
            ttk.Label(
                frame_info,
                text=f"{clave}: {valor}"
            ).pack(anchor=tk.W, padx=20)

        ttk.Label(
            frame_info,
            text="Componentes Instalados:",
            font=("Segoe UI", 11, "bold")
        ).pack(anchor=tk.W, pady=(10, 5))

        for clave, valor in componentes.items():
            label = ttk.Label(
                frame_info,
                text=f"{clave}: {valor}"
            )
            label.pack(anchor=tk.W, padx=20)

            if "No instalado" in valor:
                label.configure(foreground="#e74c3c")

        ttk.Label(
            frame_info,
            text="Modelos de Vosk:",
            font=("Segoe UI", 11, "bold")
        ).pack(anchor=tk.W, pady=(10, 5))

        ttk.Label(
            frame_info,
            text=info_modelos,
            foreground="#27ae60" if self.modelos_vosk else "#e74c3c"
        ).pack(anchor=tk.W, padx=20)

        ttk.Label(
            frame_info,
            text="Problemas Detectados:",
            font=("Segoe UI", 11, "bold")
        ).pack(anchor=tk.W, pady=(10, 5))

        problemas = []

        if not vosk_disponible:
            problemas.append("Vosk no está instalado. El modo offline no estará disponible.")

        if vosk_disponible and not self.modelos_vosk:
            problemas.append("No se encontraron modelos de Vosk. El modo offline no funcionará correctamente.")

        if not encriptacion_disponible:
            problemas.append("Cryptography no está instalado. La encriptación no estará disponible.")

        if not pydub_disponible:
            problemas.append("pydub no está instalado. La conversión de formatos no estará disponible.")

        if problemas:
            for problema in problemas:
                ttk.Label(
                    frame_info,
                    text=f"• {problema}",
                    foreground="#e74c3c"
                ).pack(anchor=tk.W, padx=20)
        else:
            ttk.Label(
                frame_info,
                text="No se detectaron problemas. La instalación es completa.",
                foreground="#27ae60"
            ).pack(anchor=tk.W, padx=20)

        frame_botones = ttk.Frame(frame_info)
        frame_botones.pack(fill=tk.X, pady=(20, 0))

        ttk.Button(
            frame_botones,
            text="Cerrar",
            command=ventana_info.destroy
        ).pack(side=tk.RIGHT)

    def mostrar_acerca_de(self):
        """Muestra información sobre la aplicación"""
        ventana_acerca = tk.Toplevel(self.root)
        ventana_acerca.title("Acerca de")
        ventana_acerca.geometry("500x400")
        ventana_acerca.resizable(False, False)

        frame_acerca = ttk.Frame(ventana_acerca, padding=20)
        frame_acerca.pack(fill=tk.BOTH, expand=True)

        ttk.Label(
            frame_acerca,
            text="Transcriptor Multilingüe con Privacidad",
            font=("Segoe UI", 16, "bold")
        ).pack(pady=(0, 5))

        ttk.Label(
            frame_acerca,
            text="Versión 1.0.0",
            font=("Segoe UI", 10)
        ).pack()

        ttk.Separator(frame_acerca, orient=tk.HORIZONTAL).pack(fill=tk.X, pady=20)

        descripcion = (
            "Esta aplicación permite la transcripción de audio en español e inglés "
            "con funciones avanzadas de privacidad para proteger información confidencial.\n\n"
            "Características principales:\n"
            "• Reconocimiento en español e inglés\n"
            "• Modo offline para privacidad total\n"
            "• Encriptación de archivos\n"
            "• Detección de palabras clave corporativas\n"
            "• Registro de actividad para auditoría\n\n"
            "Desarrollado para proteger la privacidad de información empresarial sensible."
        )

        ttk.Label(
            frame_acerca,
            text=descripcion,
            wraplength=460,
            justify=tk.LEFT
        ).pack(anchor=tk.W)

        ttk.Separator(frame_acerca, orient=tk.HORIZONTAL).pack(fill=tk.X, pady=20)

        ttk.Label(
            frame_acerca,
            text=f"© {datetime.now().year} Transcriptor Multilingüe con Privacidad"
        ).pack()

        ttk.Button(
            frame_acerca,
            text="Cerrar",
            command=ventana_acerca.destroy
        ).pack(pady=(20, 0))

    def ver_registros(self):
        """Muestra los registros de actividad"""
        dir_logs = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
        if not os.path.exists(dir_logs):
            messagebox.showinfo("Información", "No hay registros disponibles")
            return

        try:
            archivos_log = [f for f in os.listdir(dir_logs) if f.endswith('.log')]

            if not archivos_log:
                messagebox.showinfo("Información", "No hay archivos de registro disponibles")
                return

            archivos_log.sort(reverse=True)

            ventana_logs = tk.Toplevel(self.root)
            ventana_logs.title("Registros de Actividad")
            ventana_logs.geometry("700x500")
            ventana_logs.minsize(600, 400)

            frame_logs = ttk.Frame(ventana_logs, padding=10)
            frame_logs.pack(fill=tk.BOTH, expand=True)

            frame_selector = ttk.Frame(frame_logs)
            frame_selector.pack(fill=tk.X, pady=(0, 10))

            ttk.Label(frame_selector, text="Archivo de registro:").pack(side=tk.LEFT, padx=(0, 10))

            archivo_seleccionado = tk.StringVar(value=archivos_log[0])

            combo_archivos = ttk.Combobox(
                frame_selector,
                textvariable=archivo_seleccionado,
                values=archivos_log,
                state="readonly",
                width=30
            )
            combo_archivos.pack(side=tk.LEFT)

            texto_logs = tk.Text(frame_logs, wrap=tk.WORD, height=20, font=("Courier New", 10))
            texto_logs.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

            scrollbar_logs = ttk.Scrollbar(texto_logs, command=texto_logs.yview)
            scrollbar_logs.pack(side=tk.RIGHT, fill=tk.Y)
            texto_logs.config(yscrollcommand=scrollbar_logs.set)

            def cargar_log(event=None):
                archivo = archivo_seleccionado.get()
                ruta_completa = os.path.join(dir_logs, archivo)

                try:
                    with open(ruta_completa, 'r', encoding='utf-8') as f:
                        contenido = f.read()

                    texto_logs.delete("1.0", tk.END)
                    texto_logs.insert(tk.END, contenido)
                except Exception as e:
                    texto_logs.delete("1.0", tk.END)
                    texto_logs.insert(tk.END, f"Error al leer archivo: {str(e)}")

            combo_archivos.bind("<<ComboboxSelected>>", cargar_log)

            frame_botones = ttk.Frame(frame_logs)
            frame_botones.pack(fill=tk.X)

            ttk.Button(
                frame_botones,
                text="Cerrar",
                command=ventana_logs.destroy
            ).pack(side=tk.RIGHT)

            ttk.Button(
                frame_botones,
                text="Exportar",
                command=lambda: self.exportar_log(archivo_seleccionado.get())
            ).pack(side=tk.RIGHT, padx=(0, 10))

            cargar_log()

        except Exception as e:
            print(f"Error al ver registros: {e}")
            messagebox.showerror("Error", f"No se pudieron cargar los registros: {str(e)}")

    def exportar_log(self, nombre_archivo):
        """Exporta un archivo de registro"""
        if not nombre_archivo:
            return

        ruta_original = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs", nombre_archivo)

        if not os.path.exists(ruta_original):
            messagebox.showwarning("Advertencia", "El archivo no existe")
            return

        ruta_destino = filedialog.asksaveasfilename(
            title="Exportar Registro",
            defaultextension=".log",
            initialfile=nombre_archivo,
            filetypes=[
                ("Archivos de Registro", "*.log"),
                ("Archivos de Texto", "*.txt"),
                ("Todos los Archivos", "*.*")
            ]
        )

        if ruta_destino:
            try:
                shutil.copy2(ruta_original, ruta_destino)

                messagebox.showinfo("Éxito", f"Registro exportado a: {ruta_destino}")
                print(f"Registro exportado: {nombre_archivo} -> {ruta_destino}")

            except Exception as e:
                print(f"Error al exportar registro: {e}")
                messagebox.showerror("Error", f"No se pudo exportar el registro: {str(e)}")

def main():
    """Función principal para iniciar la aplicación"""
    root = tk.Tk()
    app = TranscriptorMultilingue(root)
    root.mainloop()

if __name__ == "__main__":
    root = tk.Tk()
    app = TranscriptorMultilingue(root)
    root.mainloop()