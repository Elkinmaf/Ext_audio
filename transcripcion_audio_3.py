"""
Programa de Transcripción de Audio en Tiempo Real - Español e Inglés
====================================================================
Versión compatible que no requiere webrtcvad ni sounddevice.

Requisitos:
- Python 3.7+
- Bibliotecas: SpeechRecognition, PyAudio, Numpy, Pydub (opcional)

Instalación de dependencias básicas:
```bash
pip install SpeechRecognition pyaudio numpy
```

Dependencias opcionales:
```bash
pip install pydub python-docx
```
"""

import speech_recognition as sr
import pyaudio
import wave
import os
import tkinter as tk
from tkinter import filedialog, ttk, messagebox
import threading
import time
import tempfile
import queue
import math
from datetime import datetime
import numpy as np

# Importaciones opcionales
try:
    from pydub import AudioSegment
    pydub_disponible = True
except ImportError:
    pydub_disponible = False
    print("Nota: pydub no está instalado. La conversión de formatos no estará disponible.")


class TranscripcionApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Transcriptor de Audio en Tiempo Real (Compatible)")
        self.root.geometry("800x600")
        self.root.resizable(True, True)
        
        # Configuración inicial
        self.frames = []
        self.grabando = False
        self.transcribiendo = False
        self.formato = pyaudio.paInt16
        self.canales = 1
        self.tasa_muestreo = 16000  # Óptimo para reconocimiento de voz
        self.chunk = 1024
        self.audio = pyaudio.PyAudio()
        self.reconocedor = sr.Recognizer()
        
        # Configuración optimizada para mejor captación
        self.reconocedor.energy_threshold = 300  # Umbral de energía más bajo para detectar voz
        self.reconocedor.dynamic_energy_threshold = True  # Ajuste dinámico para ambientes ruidosos
        self.reconocedor.dynamic_energy_adjustment_damping = 0.15  # Ajuste más rápido
        self.reconocedor.pause_threshold = 0.6  # Pausa más corta entre frases
        self.reconocedor.operation_timeout = None  # Sin tiempo límite
        self.reconocedor.phrase_threshold = 0.3  # Umbral para detectar frases
        self.reconocedor.non_speaking_duration = 0.3  # Duración más corta para no-voz
        
        # Ajustes para transcripción en tiempo real
        self.cola_audio = queue.Queue()
        self.duracion_segmento = 3  # Duración en segundos de cada segmento (más corto para mejor captación)
        self.superposicion = 1  # Segundos de superposición entre segmentos para no perder palabras
        self.dir_temp = tempfile.gettempdir()
        self.contador_segmentos = 0
        self.tiempo_inicio = 0
        self.duracion_total = 0
        
        # Crear la interfaz gráfica
        self.crear_interfaz()
        
    def crear_interfaz(self):
        # Frame principal
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Panel superior con estadísticas
        stats_frame = ttk.LabelFrame(main_frame, text="Información de Grabación", padding="10")
        stats_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Grid para estadísticas
        for i in range(3):
            stats_frame.columnconfigure(i, weight=1)
        
        # Tiempo de grabación
        ttk.Label(stats_frame, text="Tiempo de grabación:").grid(row=0, column=0, sticky=tk.W)
        self.tiempo_var = tk.StringVar(value="00:00:00")
        ttk.Label(stats_frame, textvariable=self.tiempo_var, font=("Arial", 10, "bold")).grid(row=0, column=1, sticky=tk.W)
        
        # Tamaño del archivo
        ttk.Label(stats_frame, text="Tamaño estimado:").grid(row=1, column=0, sticky=tk.W)
        self.tamano_var = tk.StringVar(value="0 MB")
        ttk.Label(stats_frame, textvariable=self.tamano_var).grid(row=1, column=1, sticky=tk.W)
        
        # Estado del reconocimiento
        ttk.Label(stats_frame, text="Estado:").grid(row=0, column=2, sticky=tk.W)
        self.estado_var = tk.StringVar(value="Listo")
        ttk.Label(stats_frame, textvariable=self.estado_var, font=("Arial", 10, "bold")).grid(row=0, column=3, sticky=tk.W)
        
        # Área de texto para la transcripción
        ttk.Label(main_frame, text="Transcripción en tiempo real:", font=("Arial", 11, "bold")).pack(anchor=tk.W, pady=(5, 5))
        
        text_frame = ttk.Frame(main_frame)
        text_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        self.texto_transcripcion = tk.Text(text_frame, wrap=tk.WORD, height=15, font=("Arial", 11))
        self.texto_transcripcion.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(text_frame, command=self.texto_transcripcion.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.texto_transcripcion.config(yscrollcommand=scrollbar.set)
        
        # Panel de control
        control_frame = ttk.LabelFrame(main_frame, text="Control", padding="10")
        control_frame.pack(fill=tk.X, pady=(0, 10))
        
        # Selector de idioma
        idioma_frame = ttk.Frame(control_frame)
        idioma_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(idioma_frame, text="Idioma:").pack(side=tk.LEFT, padx=(0, 10))
        self.idioma_var = tk.StringVar(value="es-ES")
        
        idiomas = [
            ("Español", "es-ES"),
            ("Inglés (US)", "en-US"),
            ("Inglés (UK)", "en-GB"),
            ("Español (MX)", "es-MX")
        ]
        
        for i, (texto, valor) in enumerate(idiomas):
            ttk.Radiobutton(
                idioma_frame, 
                text=texto, 
                value=valor, 
                variable=self.idioma_var
            ).pack(side=tk.LEFT, padx=(0, 10))
        
        # Configuraciones avanzadas
        config_frame = ttk.LabelFrame(control_frame, text="Configuración Avanzada", padding=10)
        config_frame.pack(fill=tk.X, pady=5)
        
        # Grid para configuración
        config_frame.columnconfigure(0, weight=1)
        config_frame.columnconfigure(1, weight=1)
        config_frame.columnconfigure(2, weight=1)
        config_frame.columnconfigure(3, weight=1)
        
        # Duración de segmento
        ttk.Label(config_frame, text="Duración segmento (seg):").grid(row=0, column=0, sticky=tk.W, padx=5)
        self.duracion_var = tk.StringVar(value=str(self.duracion_segmento))
        duracion_spin = ttk.Spinbox(
            config_frame,
            from_=1,
            to=10,
            width=5,
            textvariable=self.duracion_var,
            command=lambda: setattr(self, 'duracion_segmento', int(self.duracion_var.get()))
        )
        duracion_spin.grid(row=0, column=1, sticky=tk.W, padx=5)
        
        # Superposición de segmentos
        ttk.Label(config_frame, text="Superposición (seg):").grid(row=0, column=2, sticky=tk.W, padx=5)
        self.superposicion_var = tk.StringVar(value=str(self.superposicion))
        superposicion_spin = ttk.Spinbox(
            config_frame,
            from_=0.5,
            to=3,
            increment=0.5,
            width=5,
            textvariable=self.superposicion_var,
            command=lambda: setattr(self, 'superposicion', float(self.superposicion_var.get()))
        )
        superposicion_spin.grid(row=0, column=3, sticky=tk.W, padx=5)
        
        # Umbral de energía
        ttk.Label(config_frame, text="Umbral de energía:").grid(row=1, column=0, sticky=tk.W, padx=5)
        self.energia_var = tk.StringVar(value=str(self.reconocedor.energy_threshold))
        energia_spin = ttk.Spinbox(
            config_frame,
            from_=50,
            to=1000,
            increment=50,
            width=5,
            textvariable=self.energia_var,
            command=lambda: setattr(self.reconocedor, 'energy_threshold', int(self.energia_var.get()))
        )
        energia_spin.grid(row=1, column=1, sticky=tk.W, padx=5)
        
        # Umbral de pausa
        ttk.Label(config_frame, text="Umbral de pausa (seg):").grid(row=1, column=2, sticky=tk.W, padx=5)
        self.pausa_var = tk.StringVar(value=str(self.reconocedor.pause_threshold))
        pausa_spin = ttk.Spinbox(
            config_frame,
            from_=0.3,
            to=2.0,
            increment=0.1,
            width=5,
            textvariable=self.pausa_var,
            command=lambda: setattr(self.reconocedor, 'pause_threshold', float(self.pausa_var.get()))
        )
        pausa_spin.grid(row=1, column=3, sticky=tk.W, padx=5)
        
        # Opciones adicionales
        opciones_frame = ttk.Frame(config_frame)
        opciones_frame.grid(row=2, column=0, columnspan=4, sticky=tk.W, pady=5)
        
        # Opciones de procesamiento
        self.ajuste_dinamico_var = tk.BooleanVar(value=self.reconocedor.dynamic_energy_threshold)
        ttk.Checkbutton(
            opciones_frame,
            text="Ajuste dinámico",
            variable=self.ajuste_dinamico_var,
            command=lambda: setattr(self.reconocedor, 'dynamic_energy_threshold', self.ajuste_dinamico_var.get())
        ).pack(side=tk.LEFT, padx=(0, 15))
        
        self.cancelacion_ruido_var = tk.BooleanVar(value=True and pydub_disponible)
        ttk.Checkbutton(
            opciones_frame,
            text="Cancelación de ruido",
            variable=self.cancelacion_ruido_var,
            state=tk.NORMAL if pydub_disponible else tk.DISABLED
        ).pack(side=tk.LEFT, padx=(0, 15))
        
        # IMPORTANTE: Crear un frame separado para los botones directamente en main_frame
        action_frame = ttk.Frame(main_frame)
        action_frame.pack(fill=tk.X, pady=10)
        
        self.boton_grabar = ttk.Button(
            action_frame,
            text="Iniciar Grabación", 
            command=self.toggle_grabacion
        )
        self.boton_grabar.pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(
            action_frame,
            text="Cargar Archivo", 
            command=self.cargar_archivo
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(
            action_frame,
            text="Guardar Transcripción", 
            command=self.guardar_transcripcion
        ).pack(side=tk.LEFT, padx=(0, 10))
        
        ttk.Button(
            action_frame,
            text="Limpiar",
            command=self.limpiar_transcripcion
        ).pack(side=tk.LEFT)
        
        # Barra de estado en la parte inferior
        status_bar = ttk.Frame(main_frame)
        status_bar.pack(fill=tk.X, side=tk.BOTTOM)
        ttk.Label(status_bar, text="© 2025 Transcriptor en Tiempo Real (Compatible)", relief=tk.SUNKEN, anchor=tk.W).pack(fill=tk.X)
        
        # Configurar estilo personalizado
        self.configurar_estilo()
    
    def configurar_estilo(self):
        # Crear estilo personalizado para botones
        style = ttk.Style()
        if 'Accent.TButton' not in style.theme_names():
            style.configure('Accent.TButton', font=('Arial', 10, 'bold'))
            if os.name == 'nt':  # Windows
                style.map('Accent.TButton',
                    background=[('active', '!disabled', '#78909c')],
                    foreground=[('active', '!disabled', 'white')])
    
    def limpiar_transcripcion(self):
        """Limpia el área de transcripción"""
        self.texto_transcripcion.delete("1.0", tk.END)
    
    def actualizar_tiempo(self):
        # Actualizar el tiempo transcurrido mientras está grabando
        if self.grabando:
            tiempo_transcurrido = time.time() - self.tiempo_inicio
            # Formatear como HH:MM:SS
            horas = int(tiempo_transcurrido // 3600)
            minutos = int((tiempo_transcurrido % 3600) // 60)
            segundos = int(tiempo_transcurrido % 60)
            self.tiempo_var.set(f"{horas:02d}:{minutos:02d}:{segundos:02d}")
            
            # Actualizar tamaño estimado (16 bits, 16kHz, mono)
            tamano_mb = (tiempo_transcurrido * self.tasa_muestreo * 2) / (1024 * 1024)
            self.tamano_var.set(f"{tamano_mb:.2f} MB")
            
            # Programar la próxima actualización
            self.root.after(1000, self.actualizar_tiempo)
    
    def toggle_grabacion(self):
        if not self.grabando:
            # Iniciar grabación
            self.grabando = True
            self.transcribiendo = True
            self.contador_segmentos = 0
            self.tiempo_inicio = time.time()
            self.duracion_total = 0
            
            # Actualizar botón y estado
            self.boton_grabar.config(text="Detener Grabación")
            self.estado_var.set("⚫ Grabando...")
            
            # Limpiar el área de texto si es una nueva grabación
            self.texto_transcripcion.delete("1.0", tk.END)
            self.texto_transcripcion.insert(tk.END, "Escuchando... la transcripción aparecerá aquí en tiempo real.\n\n")
            
            # Iniciar temporizador
            self.actualizar_tiempo()
            
            # Iniciar grabación y transcripción en hilos separados
            threading.Thread(target=self.grabar_audio_continuo, daemon=True).start()
            threading.Thread(target=self.procesar_audio_tiempo_real, daemon=True).start()
        else:
            # Detener grabación
            self.grabando = False
            self.transcribiendo = False
            self.boton_grabar.config(text="Iniciar Grabación")
            self.estado_var.set("Procesando últimos segmentos...")
            
            # Esperar a que se procesen los últimos segmentos
            self.root.after(5000, lambda: self.estado_var.set("Listo"))
    
    def grabar_audio_continuo(self):
        """Graba audio continuamente y lo divide en segmentos para procesamiento en tiempo real"""
        stream = self.audio.open(
            format=self.formato,
            channels=self.canales,
            rate=self.tasa_muestreo,
            input=True,
            frames_per_buffer=self.chunk
        )
        
        # Inicializar variables para segmentación
        frames_buffer = []  # Buffer circular para mantener datos de varios segmentos
        max_frames_buffer = int(self.tasa_muestreo / self.chunk * (self.duracion_segmento + self.superposicion) * 2)
        frames_por_segmento = int(self.tasa_muestreo / self.chunk * self.duracion_segmento)
        contador_frames = 0
        ultimo_frame_procesado = 0
        
        try:
            # Grabar audio mientras self.grabando sea True
            while self.grabando:
                try:
                    data = stream.read(self.chunk)
                    frames_buffer.append(data)
                    contador_frames += 1
                    
                    # Limitar tamaño del buffer
                    if len(frames_buffer) > max_frames_buffer:
                        frames_buffer = frames_buffer[-max_frames_buffer:]
                    
                    # Cuando acumulamos suficientes frames nuevos, procesamos un segmento
                    # con superposición respecto al anterior
                    frames_desde_ultimo = contador_frames - ultimo_frame_procesado
                    if frames_desde_ultimo >= frames_por_segmento - int(self.tasa_muestreo / self.chunk * self.superposicion):
                        ultimo_frame_procesado = contador_frames
                        # Tomar los últimos frames_por_segmento del buffer
                        frames_segmento = frames_buffer[-frames_por_segmento:]
                        self.procesar_segmento(frames_segmento)
                
                except Exception as e:
                    print(f"Error durante la grabación: {e}")
            
            # Procesar el último segmento parcial si existe
            if frames_buffer and contador_frames > ultimo_frame_procesado:
                frames_segmento = frames_buffer[-min(len(frames_buffer), frames_por_segmento):]
                self.procesar_segmento(frames_segmento)
        
        finally:
            # Cerrar el stream
            stream.stop_stream()
            stream.close()
    
    def procesar_segmento(self, frames_segmento):
        """Guarda un segmento de audio y lo pone en la cola para transcripción"""
        try:
            # Generar nombre único para el archivo temporal
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            nombre_segmento = f"seg_{timestamp}_{self.contador_segmentos}.wav"
            ruta_segmento = os.path.join(self.dir_temp, nombre_segmento)
            self.contador_segmentos += 1
            
            # Guardar el segmento como archivo WAV
            with wave.open(ruta_segmento, 'wb') as wf:
                wf.setnchannels(self.canales)
                wf.setsampwidth(self.audio.get_sample_size(self.formato))
                wf.setframerate(self.tasa_muestreo)
                wf.writeframes(b''.join(frames_segmento))
            
            # Preprocesar el audio si la cancelación de ruido está activada
            if self.cancelacion_ruido_var.get() and pydub_disponible:
                ruta_procesada = self.preprocesar_audio(ruta_segmento)
                if ruta_procesada:
                    # Si el procesamiento fue exitoso, usamos el archivo procesado
                    try:
                        os.remove(ruta_segmento)  # Eliminar archivo original
                    except:
                        pass
                    ruta_segmento = ruta_procesada
            
            # Añadir el segmento a la cola para transcripción
            self.cola_audio.put(ruta_segmento)
        
        except Exception as e:
            print(f"Error al procesar segmento: {e}")
    
    def preprocesar_audio(self, ruta_archivo):
        """Aplica preprocesamiento al audio para mejorar la calidad"""
        if not pydub_disponible:
            return ruta_archivo
        
        try:
            # Cargar el audio con pydub
            audio = AudioSegment.from_wav(ruta_archivo)
            
            # Normalizar volumen (aumentar volumen bajo, reducir volumen alto)
            audio = audio.normalize()
            
            # Aplicar filtro de paso alto para eliminar ruido de baja frecuencia
            audio = audio.high_pass_filter(80)
            
            # Generar nombre para archivo procesado
            ruta_procesada = ruta_archivo.replace(".wav", "_proc.wav")
            
            # Exportar audio procesado
            audio.export(ruta_procesada, format="wav")
            
            return ruta_procesada
        
        except Exception as e:
            print(f"Error en preprocesamiento: {e}")
            return ruta_archivo  # En caso de error, devolver ruta original
    
    def procesar_audio_tiempo_real(self):
        """Procesa los segmentos de audio de la cola para transcripción en tiempo real"""
        try:
            # Guardar transcripciones anteriores para contexto
            transcripciones_previas = []
            
            while self.transcribiendo or not self.cola_audio.empty():
                try:
                    # Esperar hasta que haya un segmento disponible
                    if not self.cola_audio.empty():
                        archivo_segmento = self.cola_audio.get(block=False)
                        
                        # Transcribir segmento
                        texto = self.transcribir_segmento(archivo_segmento, transcripciones_previas)
                        
                        # Guardar transcripción para contexto si tiene contenido
                        if texto and texto.strip():
                            transcripciones_previas.append(texto)
                            # Mantener solo las últimas 5 transcripciones como contexto
                            if len(transcripciones_previas) > 5:
                                transcripciones_previas = transcripciones_previas[-5:]
                        
                        # Eliminar archivo temporal
                        try:
                            os.remove(archivo_segmento)
                        except:
                            pass
                    else:
                        # Esperar un poco si no hay segmentos
                        time.sleep(0.2)
                        
                except queue.Empty:
                    # No hay elementos en la cola, esperar
                    time.sleep(0.2)
                    
        except Exception as e:
            print(f"Error en el procesamiento de audio en tiempo real: {e}")
            self.estado_var.set(f"Error: {str(e)}")
    
    def transcribir_segmento(self, archivo_audio, contexto_previo=None):
        """Transcribe un segmento de audio y actualiza la UI con el resultado"""
        try:
            with sr.AudioFile(archivo_audio) as fuente:
                # Ajustar el reconocedor para audio en ambiente ruidoso
                self.reconocedor.adjust_for_ambient_noise(fuente, duration=0.2)
                # Obtener audio
                audio_data = self.reconocedor.record(fuente)
                
                # Actualizar estado
                self.root.after(0, lambda: self.estado_var.set("⚫ Transcribiendo..."))
                
                # Obtener el idioma seleccionado
                idioma = self.idioma_var.get()
                
                # Transcribir el audio con contexto
                # Usar contexto previo para mejorar la precisión
                texto = ""
                
                try:
                    # Intentar usar la API de Google con reconocimiento de contexto
                    frases_contexto = []
                    if contexto_previo and len(contexto_previo) > 0:
                        # Usar últimas 3 frases como contexto
                        frases_contexto = ' '.join(contexto_previo[-3:])
                        
                    # Hacer reconocimiento
                    texto = self.reconocedor.recognize_google(
                        audio_data, 
                        language=idioma,
                        show_all=False  # Solo queremos el resultado más probable
                    )
                    
                    # Si tenemos texto y tenemos contexto previo, intentamos hacer correcciones
                    if texto and frases_contexto:
                        # Verificar si hay palabras del contexto que se repiten al inicio
                        # para evitar duplicación de frases entre segmentos
                        palabras_texto = texto.split()
                        palabras_contexto = frases_contexto.split()
                        
                        # Si tenemos pocas palabras, no hacemos corrección
                        if len(palabras_texto) > 3 and len(palabras_contexto) > 3:
                            # Buscar duplicaciones al inicio (solapamiento entre segmentos)
                            n_palabras = min(5, len(palabras_texto), len(palabras_contexto))
                            
                            # Calcular si hay solapamiento significativo
                            inicio_texto = ' '.join(palabras_texto[:n_palabras]).lower()
                            fin_contexto = ' '.join(palabras_contexto[-n_palabras:]).lower()
                            
                            similitud = self.calcular_similitud(inicio_texto, fin_contexto)
                            
                            # Si hay alta similitud, eliminar la parte duplicada
                            if similitud > 0.7:
                                texto = ' '.join(palabras_texto[n_palabras:])
                
                except sr.UnknownValueError:
                    texto = ""
                
                except Exception as e:
                    print(f"Error en reconocimiento: {e}")
                    texto = ""
                
                # Si se obtuvo texto, mostrarlo
                if texto and texto.strip():
                    # Aplicar correcciones finales
                    texto = self.aplicar_correcciones_post(texto, idioma)
                    
                    # Mostrar el texto transcrito en la interfaz
                    self.root.after(0, lambda: self.actualizar_transcripcion(texto))
                    
                # Actualizar estado
                self.root.after(0, lambda: self.estado_var.set("⚫ Grabando..."))
                
                return texto
                
        except sr.UnknownValueError:
            # Si no se reconoce nada, no mostrar error
            return ""
            
        except sr.RequestError as e:
            # Error con el servicio
            self.root.after(0, lambda: self.estado_var.set(f"Error en el servicio: {e}"))
            return ""
            
        except Exception as e:
            print(f"Error en transcripción: {e}")
            self.root.after(0, lambda: self.estado_var.set(f"Error: {str(e)}"))
            return ""
    
    def calcular_similitud(self, texto1, texto2):
        """Calcula la similitud entre dos textos (0-1)"""
        if not texto1 or not texto2:
            return 0
            
        # Método simple: contar palabras compartidas
        palabras1 = set(texto1.lower().split())
        palabras2 = set(texto2.lower().split())
        
        # Calcular coeficiente de Jaccard
        if not palabras1 or not palabras2:
            return 0
            
        interseccion = len(palabras1.intersection(palabras2))
        union = len(palabras1.union(palabras2))
        
        return interseccion / union if union > 0 else 0
    
    def aplicar_correcciones_post(self, texto, idioma):
        """Aplica correcciones post-procesamiento al texto reconocido"""
        if not texto:
            return texto
            
        # Correcciones específicas del idioma
        if idioma.startswith('es'):
            # Corregir espacios antes de signos de puntuación en español
            for signo in ['?', '!', ',', '.', ':', ';']:
                texto = texto.replace(f' {signo}', signo)
            
            # Corregir espacios en signos de interrogación/exclamación en español
            texto = texto.replace('¿ ', '¿').replace('¡ ', '¡')
            
        # Correcciones generales
        # Eliminar espacios extras
        texto = ' '.join(texto.split())
        
        # Asegurar que la primera letra esté en mayúscula
        if texto and len(texto) > 0:
            texto = texto[0].upper() + texto[1:]
            
        return texto
    
    def cargar_archivo(self):
        # Permitir al usuario seleccionar un archivo de audio
        archivo = filedialog.askopenfilename(
            title="Seleccionar Archivo de Audio",
            filetypes=[
                ("Archivos WAV", "*.wav"),
                ("Todos los Archivos", "*.*") if not pydub_disponible else 
                ("Todos los formatos compatibles", "*.wav;*.mp3;*.ogg;*.flac")
            ]
        )
        
        if archivo:
            self.estado_var.set(f"Procesando archivo: {os.path.basename(archivo)}...")
            self.texto_transcripcion.delete("1.0", tk.END)
            self.texto_transcripcion.insert(tk.END, f"Procesando archivo: {os.path.basename(archivo)}...\n\n")
            
            # Convertir a WAV si es necesario y si pydub está disponible
            extension = os.path.splitext(archivo)[1].lower()
            
            if extension != '.wav':
                if pydub_disponible:
                    # Intentar convertir si pydub está disponible
                    try:
                        self.estado_var.set("Convirtiendo archivo a formato WAV...")
                        archivo_wav = self.convertir_a_wav(archivo)
                        if archivo_wav:
                            archivo = archivo_wav
                        else:
                            self.root.after(0, lambda: self.estado_var.set("No se pudo convertir el archivo"))
                            messagebox.showwarning("Error de conversión", 
                                                "No se pudo convertir el archivo a formato WAV.")
                            return
                    except Exception as e:
                        messagebox.showwarning("Error de conversión", 
                                            f"No se pudo convertir el archivo: {str(e)}\nSe intentará procesar el archivo original.")
                else:
                    self.root.after(0, lambda: self.estado_var.set("Solo se soportan archivos WAV"))
                    messagebox.showwarning("Formato no soportado", 
                                        "Solo se soportan archivos WAV. Instale pydub para soporte de otros formatos.")
                    return
            
            # Transcribir el archivo en un hilo separado
            threading.Thread(target=lambda: self.transcribir_archivo_grande(archivo), daemon=True).start()
    
    def convertir_a_wav(self, archivo_origen):
        """Convierte un archivo de audio a formato WAV"""
        if not pydub_disponible:
            return None
            
        try:
            # Generar nombre para archivo WAV
            nombre_base = os.path.splitext(os.path.basename(archivo_origen))[0]
            archivo_destino = os.path.join(self.dir_temp, f"{nombre_base}_conv.wav")
            
            # Convertir usando pydub
            extension = os.path.splitext(archivo_origen)[1].lower()
            
            # Cargar el archivo según su extensión
            if extension == '.mp3':
                audio = AudioSegment.from_mp3(archivo_origen)
            elif extension == '.ogg':
                audio = AudioSegment.from_ogg(archivo_origen)
            elif extension == '.flac':
                audio = AudioSegment.from_file(archivo_origen, "flac")
            else:
                audio = AudioSegment.from_file(archivo_origen)
            
            # Configurar para reconocimiento de voz óptimo
            audio = audio.set_channels(1).set_frame_rate(16000)
            
            # Normalizar y aplicar filtros si está activada la opción
            if self.cancelacion_ruido_var.get():
                audio = audio.normalize()
                audio = audio.high_pass_filter(80)
            
            # Exportar a WAV
            audio.export(archivo_destino, format="wav")
            
            return archivo_destino
            
        except Exception as e:
            print(f"Error en conversión: {e}")
            return None
    
    def transcribir_archivo_grande(self, archivo_audio):
        """Transcribe un archivo de audio grande dividiéndolo en segmentos"""
        try:
            # Determinar si es un archivo WAV
            extension = os.path.splitext(archivo_audio)[1].lower()
            
            if extension != '.wav':
                if pydub_disponible:
                    # Intentar convertir si pydub está disponible
                    archivo_wav = self.convertir_a_wav(archivo_audio)
                    if archivo_wav:
                        archivo_audio = archivo_wav
                    else:
                        self.root.after(0, lambda: self.estado_var.set("No se pudo convertir el archivo"))
                        messagebox.showwarning("Error de formato", 
                                              "No se pudo convertir el archivo. Asegúrese de tener instaladas las dependencias necesarias.")
                        return
                else:
                    self.root.after(0, lambda: self.estado_var.set("Solo se soportan archivos WAV"))
                    messagebox.showwarning("Formato no soportado", 
                                          "Solo se soportan archivos WAV. Instale pydub para soporte de otros formatos.")
                    return
            
            # Abrir el archivo para determinar su duración
            with wave.open(archivo_audio, 'rb') as wf:
                # Verificar parámetros del audio
                canales = wf.getnchannels()
                tasa = wf.getframerate()
                
                # Calcular duración en segundos
                frames = wf.getnframes()
                rate = wf.getframerate()
                duracion = frames / float(rate)
                
                # Mostrar información del archivo
                info_audio = (f"Información del archivo:\n"
                             f"- Duración: {int(duracion // 60)}:{int(duracion % 60):02d}\n"
                             f"- Canales: {canales}\n"
                             f"- Tasa de muestreo: {tasa} Hz\n"
                             f"- Tamaño: {os.path.getsize(archivo_audio) / (1024*1024):.2f} MB\n\n")
                
                self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, info_audio))
                
                # Determinar la estrategia de procesamiento según la duración
                if duracion > 60:  # Más de un minuto
                    self.root.after(0, lambda: self.estado_var.set("Procesando archivo grande por segmentos..."))
                    self.procesar_archivo_por_segmentos(archivo_audio, duracion)
                else:
                    # Archivo pequeño, procesarlo normalmente
                    self.transcribir_segmento_archivo(archivo_audio)
            
        except Exception as e:
            self.root.after(0, lambda: self.estado_var.set(f"Error al procesar archivo: {str(e)}"))
            messagebox.showerror("Error", f"Error al procesar archivo: {str(e)}")
    
    def procesar_archivo_por_segmentos(self, archivo_audio, duracion_total):
        """Procesa un archivo de audio grande por segmentos con solapamiento para mejor captación"""
        try:
            # Obtener el idioma seleccionado
            idioma = self.idioma_var.get()
            
            # Tamaño de segmento en segundos (más corto para mejor precisión)
            tamano_segmento = 15  # 15 segundos por segmento
            solapamiento = 2      # 2 segundos de solapamiento entre segmentos
            
            # Calcular número de segmentos con solapamiento
            num_segmentos = math.ceil((duracion_total - solapamiento) / (tamano_segmento - solapamiento))
            
            # Actualizar interfaz
            self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, 
                            f"Archivo dividido en {num_segmentos} segmentos con solapamiento para mejor captación...\n\n"))
            
            # Lista para almacenar transcripciones previas para contexto
            transcripciones_previas = []
            
            # Utilizar el reconocedor para cada segmento
            with sr.AudioFile(archivo_audio) as fuente:
                # Ajustar para ruido ambiental al inicio del archivo
                self.reconocedor.adjust_for_ambient_noise(fuente, duration=min(1.0, duracion_total/10))
                
                for i in range(num_segmentos):
                    # Actualizar progreso
                    progreso = int((i+1) * 100 / num_segmentos)
                    self.root.after(0, lambda i=i, p=progreso: 
                                   self.estado_var.set(f"Procesando segmento {i+1}/{num_segmentos} ({p}%)..."))
                    
                    # Calcular posición de inicio del segmento con solapamiento
                    offset = i * (tamano_segmento - solapamiento)
                    
                    # Último segmento podría ser más corto
                    duracion = min(tamano_segmento, duracion_total - offset)
                    
                    if duracion <= 0:
                        break
                    
                    # Posicionar y grabar el segmento
                    fuente.rewind()
                    audio_segmento = self.reconocedor.record(fuente, duration=duracion, offset=offset)
                    
                    try:
                        # Transcribir segmento
                        texto = self.reconocedor.recognize_google(audio_segmento, language=idioma)
                        
                        # Aplicar correcciones a la transcripción
                        if texto and texto.strip():
                            # Aplicar correcciones de continuidad con segmentos previos
                            if transcripciones_previas:
                                # Verificar solapamiento con transcripción anterior
                                palabras_texto = texto.split()
                                ultima_transcripcion = transcripciones_previas[-1]
                                palabras_previas = ultima_transcripcion.split()
                                
                                # Verificar inicio para evitar duplicaciones
                                if len(palabras_texto) > 3 and len(palabras_previas) > 3:
                                    n_palabras = min(4, len(palabras_texto), len(palabras_previas))
                                    
                                    inicio_texto = ' '.join(palabras_texto[:n_palabras]).lower()
                                    fin_previo = ' '.join(palabras_previas[-n_palabras:]).lower()
                                    
                                    similitud = self.calcular_similitud(inicio_texto, fin_previo)
                                    
                                    # Si hay alta similitud, eliminar parte duplicada
                                    if similitud > 0.5:
                                        texto = ' '.join(palabras_texto[n_palabras:])
                            
                            # Aplicar correcciones según el idioma
                            texto = self.aplicar_correcciones_post(texto, idioma)
                            
                            # Guardar para contexto
                            transcripciones_previas.append(texto)
                            if len(transcripciones_previas) > 5:
                                transcripciones_previas = transcripciones_previas[-5:]
                            
                            # Mostrar el segmento de tiempo
                            horas_inicio = int(offset // 3600)
                            min_inicio = int((offset % 3600) // 60)
                            seg_inicio = int(offset % 60)
                            
                            tiempo_formateado = f"[{horas_inicio:02d}:{min_inicio:02d}:{seg_inicio:02d}]"
                            
                            # Actualizar transcripción
                            self.root.after(0, lambda tiempo=tiempo_formateado, texto=texto: 
                                          self.actualizar_transcripcion_con_tiempo(tiempo, texto))
                    
                    except sr.UnknownValueError:
                        # No se reconoció nada en este segmento
                        pass
                    
                    except Exception as e:
                        print(f"Error en segmento {i+1}: {e}")
            
            # Transcripción completada
            self.root.after(0, lambda: self.estado_var.set("Transcripción completada"))
            self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, 
                           "\n--- Fin de la transcripción ---\n"))
            
        except Exception as e:
            self.root.after(0, lambda: self.estado_var.set(f"Error: {str(e)}"))
            print(f"Error al procesar archivo por segmentos: {e}")
    
    def transcribir_segmento_archivo(self, archivo_audio):
        """Transcribe un archivo de audio completo usando técnicas optimizadas"""
        try:
            # Obtener el idioma seleccionado
            idioma = self.idioma_var.get()
            
            with sr.AudioFile(archivo_audio) as fuente:
                # Ajustar para ruido ambiental
                self.reconocedor.adjust_for_ambient_noise(fuente, duration=min(1.0, 0.1))
                
                # Grabar todo el audio
                audio_data = self.reconocedor.record(fuente)
                
                # Transcribir con la configuración actual
                texto = self.reconocedor.recognize_google(audio_data, language=idioma)
                
                if texto.strip():
                    # Aplicar correcciones post-procesamiento
                    texto = self.aplicar_correcciones_post(texto, idioma)
                    
                    # Mostrar texto transcrito
                    self.root.after(0, lambda: self.actualizar_transcripcion(texto))
                
                # Actualizar estado
                self.root.after(0, lambda: self.estado_var.set("Transcripción completada"))
        
        except sr.UnknownValueError:
            self.root.after(0, lambda: self.estado_var.set("No se pudo reconocer el audio"))
            self.root.after(0, lambda: self.texto_transcripcion.insert(tk.END, 
                           "No se pudo reconocer ningún texto en el audio.\n"
                           "Intente ajustar el umbral de energía o usar cancelación de ruido.\n"))
        
        except Exception as e:
            self.root.after(0, lambda: self.estado_var.set(f"Error: {str(e)}"))
            print(f"Error al transcribir archivo: {e}")
    
    def actualizar_transcripcion(self, texto):
        """Actualiza el área de texto con la transcripción"""
        # Insertar el texto transcrito en el área de texto
        self.texto_transcripcion.insert(tk.END, texto + "\n\n")
        self.texto_transcripcion.see(tk.END)  # Desplazar al final
    
    def actualizar_transcripcion_con_tiempo(self, tiempo, texto):
        """Actualiza el área de texto con la transcripción y marca de tiempo"""
        # Insertar el texto transcrito con marca de tiempo
        self.texto_transcripcion.insert(tk.END, f"{tiempo} {texto}\n\n")
        self.texto_transcripcion.see(tk.END)  # Desplazar al final
    
    def guardar_transcripcion(self):
        """Guarda la transcripción actual en un archivo de texto"""
        # Obtener el texto de la transcripción
        texto = self.texto_transcripcion.get("1.0", tk.END).strip()
        
        if not texto:
            messagebox.showwarning("Advertencia", "No hay texto para guardar")
            return
        
        # Obtener fecha y hora actual para nombre de archivo por defecto
        fecha_hora = datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_defecto = f"transcripcion_{fecha_hora}.txt"
        
        # Permitir al usuario seleccionar dónde guardar el archivo
        archivo = filedialog.asksaveasfilename(
            title="Guardar Transcripción",
            defaultextension=".txt",
            initialfile=nombre_defecto,
            filetypes=[
                ("Archivos de Texto", "*.txt"),
                ("Documentos Word", "*.docx") if 'docx' in globals() else None,
                ("Documentos SRT", "*.srt"),
                ("Todos los Archivos", "*.*")
            ]
        )
        
        if archivo:
            try:
                # Determinar formato de salida
                extension = os.path.splitext(archivo)[1].lower()
                
                if extension == '.docx':
                    try:
                        # Intentar usar python-docx si está instalado
                        import docx
                        
                        # Crear nuevo documento
                        doc = docx.Document()
                        
                        # Agregar título
                        doc.add_heading('Transcripción de Audio', 0)
                        
                        # Agregar fecha y hora
                        doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
                        doc.add_paragraph(f'Idioma: {self.idioma_var.get()}')
                        
                        # Agregar contenido de la transcripción
                        doc.add_paragraph(texto)
                        
                        # Guardar documento
                        doc.save(archivo)
                        
                    except ImportError:
                        messagebox.showinfo("Información", 
                                           "Librería python-docx no instalada. Guardando como texto plano.")
                        with open(archivo, 'w', encoding='utf-8') as f:
                            f.write(texto)
                
                elif extension == '.srt':
                    # Crear archivo de subtítulos SRT
                    # Primero, extraer todos los segmentos con marcas de tiempo
                    lineas = texto.split('\n')
                    subtitulos = []
                    
                    for i, linea in enumerate(lineas):
                        if linea.startswith('[') and ']' in linea:
                            # Extraer marca de tiempo y texto
                            partes = linea.split(']', 1)
                            if len(partes) == 2:
                                tiempo_str = partes[0][1:]  # Quitar corchete inicial
                                texto_subtitulo = partes[1].strip()
                                
                                # Convertir marca de tiempo a formato SRT (HH:MM:SS,mmm)
                                h, m, s = tiempo_str.split(':')
                                tiempo_inicio = f"{h}:{m}:{s},000"
                                
                                # Calcular tiempo de fin (5 segundos después)
                                h_fin, m_fin, s_fin = int(h), int(m), int(s) + 5
                                if s_fin >= 60:
                                    s_fin -= 60
                                    m_fin += 1
                                if m_fin >= 60:
                                    m_fin -= 60
                                    h_fin += 1
                                
                                tiempo_fin = f"{h_fin:02d}:{m_fin:02d}:{s_fin:02d},000"
                                
                                # Añadir a lista de subtítulos
                                subtitulos.append({
                                    'num': len(subtitulos) + 1,
                                    'inicio': tiempo_inicio,
                                    'fin': tiempo_fin,
                                    'texto': texto_subtitulo
                                })
                    
                    # Escribir archivo SRT
                    with open(archivo, 'w', encoding='utf-8') as f:
                        for sub in subtitulos:
                            f.write(f"{sub['num']}\n")
                            f.write(f"{sub['inicio']} --> {sub['fin']}\n")
                            f.write(f"{sub['texto']}\n\n")
                
                else:
                    # Guardar como texto plano
                    with open(archivo, 'w', encoding='utf-8') as f:
                        f.write(texto)
                
                messagebox.showinfo("Éxito", f"Transcripción guardada en: {archivo}")
            
            except Exception as e:
                messagebox.showerror("Error", f"Error al guardar el archivo: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = TranscripcionApp(root)
    root.mainloop()